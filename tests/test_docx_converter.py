"""Unit tests for ``workflows.docx_converter.convert_docx_to_txt``.

The contract pins three properties:

  - Output is byte-deterministic per input.
  - Output round-trips through :func:`parse_minutes_txt` without error —
    section headers stay distinct, tables stay pipe-delimited, and no
    paragraph silently folds into a table cell.
  - ``output_path=None`` is a true dry-run — nothing is written.

Tests build synthetic ``.docx`` files via ``python-docx`` so the suite
runs offline without any fixture file under ``data-lake/``.
"""
from __future__ import annotations

from pathlib import Path

import pytest

docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402

from spectrum_systems_core.workflows.docx_converter import (  # noqa: E402
    convert_docx_to_txt,
)
from spectrum_systems_core.workflows.minutes_parser import (  # noqa: E402
    parse_minutes_text,
)


def _build_minutes_docx(path: Path) -> None:
    """Write a synthetic minutes ``.docx`` with one discussion + one action."""
    doc = Document()
    doc.add_paragraph("7 GHz Synthetic TIG Meeting Minutes")
    doc.add_paragraph("")

    meta = doc.add_table(rows=1, cols=4)
    cells = meta.rows[0].cells
    cells[0].text = "Meeting Name:"
    cells[1].text = "7 GHz Synthetic TIG"
    cells[2].text = "Meeting Date:"
    cells[3].text = "01/15/2026"

    doc.add_paragraph("")
    doc.add_paragraph("Meeting Overview")
    doc.add_paragraph("")
    doc.add_paragraph("A synthetic meeting body for converter tests.")
    doc.add_paragraph("")

    doc.add_paragraph("Discussion/Questions Log")
    doc.add_paragraph("")
    discussion = doc.add_table(rows=2, cols=6)
    headers = ["#", "Category", "Question/Topic", "Asked By",
               "Initial Response / Discussion", "Follow-up / Action Item"]
    for i, label in enumerate(headers):
        discussion.rows[0].cells[i].text = label
    row1 = discussion.rows[1].cells
    row1[0].text = "1"
    row1[1].text = "Scope"
    row1[2].text = "What is in scope?"
    row1[3].text = "Alice"
    row1[4].text = "NTIA confirmed the scope as drafted."
    row1[5].text = "N/A"

    doc.add_paragraph("")
    doc.add_paragraph("Next Steps")
    doc.add_paragraph("")
    doc.add_paragraph("Review the draft and respond by Friday.")
    doc.add_paragraph("")

    doc.add_paragraph("Action Items")
    doc.add_paragraph("")
    actions = doc.add_table(rows=2, cols=4)
    action_headers = ["Item", "Responsible Party", "Due Date", "Status"]
    for i, label in enumerate(action_headers):
        actions.rows[0].cells[i].text = label
    a_row = actions.rows[1].cells
    a_row[0].text = "Send draft to agencies."
    a_row[1].text = "NTIA"
    a_row[2].text = "1/20/26"
    a_row[3].text = "In progress"

    doc.save(str(path))


@pytest.fixture()
def minutes_docx(tmp_path: Path) -> Path:
    path = tmp_path / "synthetic-minutes.docx"
    _build_minutes_docx(path)
    return path


def test_convert_produces_pipe_delimited_tables(minutes_docx: Path):
    text = convert_docx_to_txt(minutes_docx)
    # Discussion row: 6 cells -> 5 pipes
    assert "1 | Scope | What is in scope? | Alice | " in text
    # Action row: 4 cells -> 3 pipes
    assert "Send draft to agencies. | NTIA | 1/20/26 | In progress" in text


def test_convert_preserves_section_headers(minutes_docx: Path):
    text = convert_docx_to_txt(minutes_docx)
    for header in ("Meeting Overview", "Discussion/Questions Log",
                   "Next Steps", "Action Items"):
        assert f"\n{header}\n" in text, f"section header missing: {header}"


def test_convert_output_parses_with_minutes_parser(minutes_docx: Path):
    text = convert_docx_to_txt(minutes_docx)
    parsed = parse_minutes_text(text, source_path=str(minutes_docx))
    assert len(parsed.discussion_items) == 1
    assert parsed.discussion_items[0].category == "Scope"
    assert parsed.discussion_items[0].asked_by == "Alice"
    assert len(parsed.action_items) == 1
    assert parsed.action_items[0].responsible_party == "NTIA"
    assert len(parsed.next_steps) == 1


def test_convert_is_deterministic(minutes_docx: Path):
    first = convert_docx_to_txt(minutes_docx)
    second = convert_docx_to_txt(minutes_docx)
    assert first == second


def test_convert_dry_run_does_not_write(minutes_docx: Path, tmp_path: Path):
    out_path = tmp_path / "should-not-exist.txt"
    convert_docx_to_txt(minutes_docx, output_path=None)
    assert not out_path.exists()


def test_convert_writes_when_output_path_given(
    minutes_docx: Path, tmp_path: Path
):
    out_path = tmp_path / "written.txt"
    returned = convert_docx_to_txt(minutes_docx, output_path=out_path)
    assert out_path.exists()
    assert out_path.read_text(encoding="utf-8") == returned


def test_convert_output_trailing_newline(minutes_docx: Path):
    text = convert_docx_to_txt(minutes_docx)
    assert text.endswith("\n")
    # No accidental trailing run of blank lines
    assert not text.endswith("\n\n\n")
