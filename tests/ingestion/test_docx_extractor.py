"""Tests for DocxExtractor (Phase L.0).

All .docx fixtures are built in-memory using python-docx's Document() class.
No binary fixtures, no mocks (except test_extract_never_raises which
monkeypatches Document to verify the never-raises contract).
"""
from __future__ import annotations

import io
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from spectrum_systems_core.cli import extract_docx
from spectrum_systems_core.ingestion.docx_extractor import DocxExtractor


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(str(path))


def _write_docx_with_table(path: Path, table_rows: list[list[str]]) -> None:
    """Write a .docx with a single table whose contents are ``table_rows``."""
    doc = Document()
    if not table_rows:
        doc.save(str(path))
        return
    cols = max(len(r) for r in table_rows)
    table = doc.add_table(rows=len(table_rows), cols=cols)
    for r_idx, row in enumerate(table_rows):
        for c_idx, cell_text in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = cell_text
    doc.save(str(path))


class TestExtractSingleDocx(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_extract_single_docx_produces_txt(self) -> None:
        src = self.tmp / "meeting.docx"
        _write_docx(src, ["Hello world", "Second paragraph"])
        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        txt = Path(result["output_path"])
        self.assertTrue(txt.is_file())
        self.assertTrue(txt.suffix == ".txt")

    def test_output_path_defaults_to_same_dir_as_input(self) -> None:
        src = self.tmp / "notes.docx"
        _write_docx(src, ["Content"])
        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        expected = self.tmp / "notes.txt"
        self.assertEqual(Path(result["output_path"]), expected)
        self.assertTrue(expected.is_file())

    def test_output_path_custom_dir(self) -> None:
        src = self.tmp / "doc.docx"
        out_dir = self.tmp / "out"
        _write_docx(src, ["Text"])
        result = DocxExtractor().extract(str(src), output_path=str(out_dir / "doc.txt"))
        self.assertEqual(result["status"], "success")
        self.assertTrue((out_dir / "doc.txt").is_file())

    def test_paragraphs_joined_with_double_newline(self) -> None:
        src = self.tmp / "multi.docx"
        _write_docx(src, ["First", "Second", "Third"])
        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        content = Path(result["output_path"]).read_text(encoding="utf-8")
        self.assertEqual(content, "First\n\nSecond\n\nThird")

    def test_empty_paragraphs_skipped(self) -> None:
        src = self.tmp / "sparse.docx"
        _write_docx(src, ["", "Real content", "", "More content", ""])
        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        content = Path(result["output_path"]).read_text(encoding="utf-8")
        self.assertEqual(content, "Real content\n\nMore content")
        self.assertEqual(result["paragraph_count"], 2)

    def test_missing_file_returns_failure_not_exception(self) -> None:
        result = DocxExtractor().extract("/nonexistent/path/meeting.docx")
        self.assertEqual(result["status"], "failure")
        self.assertIn("file_not_found", result["reason"])
        self.assertIsNone(result["output_path"])
        self.assertEqual(result["paragraph_count"], 0)
        self.assertEqual(result["character_count"], 0)

    def test_empty_document_returns_failure(self) -> None:
        src = self.tmp / "empty.docx"
        doc = Document()
        doc.save(str(src))
        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "failure")
        self.assertIn("empty_document", result["reason"])
        self.assertIsNone(result["output_path"])

    def test_paragraph_count_and_character_count_populated(self) -> None:
        src = self.tmp / "counted.docx"
        _write_docx(src, ["Hello", "World"])
        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        self.assertEqual(result["paragraph_count"], 2)
        expected_chars = len("Hello\n\nWorld")
        self.assertEqual(result["character_count"], expected_chars)

    def test_no_partial_write_on_empty_document(self) -> None:
        src = self.tmp / "empty2.docx"
        doc = Document()
        doc.save(str(src))
        DocxExtractor().extract(str(src))
        txt = self.tmp / "empty2.txt"
        self.assertFalse(txt.exists(), "must not write .txt when extraction fails")


class TestExtractTables(unittest.TestCase):
    """Tables-in-document-order extraction (fix for hollow meeting minutes)."""

    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_table_content_extracted(self) -> None:
        src = self.tmp / "minutes.docx"
        _write_docx_with_table(
            src,
            [
                ["Topic", "Owner", "Notes"],
                ["Kickoff", "Alice", "Reviewed agenda"],
                ["Action items", "Bob", "Send draft by Friday"],
            ],
        )
        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        content = Path(result["output_path"]).read_text(encoding="utf-8")
        # Every cell value must appear in the output.
        for cell in [
            "Topic", "Owner", "Notes",
            "Kickoff", "Alice", "Reviewed agenda",
            "Action items", "Bob", "Send draft by Friday",
        ]:
            self.assertIn(cell, content)
        # Cells inside a row are joined with ' | '.
        self.assertIn("Kickoff | Alice | Reviewed agenda", content)
        self.assertEqual(result["table_count"], 1)
        self.assertEqual(result["table_row_count"], 3)

    def test_paragraph_and_table_interleaved(self) -> None:
        src = self.tmp / "interleaved.docx"
        doc = Document()
        doc.add_paragraph("Section 1: Intro")
        t1 = doc.add_table(rows=1, cols=2)
        t1.rows[0].cells[0].text = "T1-A"
        t1.rows[0].cells[1].text = "T1-B"
        doc.add_paragraph("Section 2: Discussion")
        t2 = doc.add_table(rows=1, cols=2)
        t2.rows[0].cells[0].text = "T2-A"
        t2.rows[0].cells[1].text = "T2-B"
        doc.add_paragraph("Section 3: Conclusion")
        doc.save(str(src))

        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        content = Path(result["output_path"]).read_text(encoding="utf-8")
        # Document order must be preserved.
        positions = [
            content.find("Section 1: Intro"),
            content.find("T1-A | T1-B"),
            content.find("Section 2: Discussion"),
            content.find("T2-A | T2-B"),
            content.find("Section 3: Conclusion"),
        ]
        self.assertTrue(all(p >= 0 for p in positions), positions)
        self.assertEqual(positions, sorted(positions))
        self.assertEqual(result["table_count"], 2)
        self.assertEqual(result["table_row_count"], 2)

    def test_empty_table_cells_skipped(self) -> None:
        src = self.tmp / "sparse_table.docx"
        _write_docx_with_table(
            src,
            [
                ["A", "", "C"],
                ["", "", ""],   # entirely empty row — must not emit blank line
                ["D", "", ""],
            ],
        )
        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        content = Path(result["output_path"]).read_text(encoding="utf-8")
        # No empty rows produce blank "  | " artifacts.
        self.assertNotIn("\n\n\n", content)
        # Rows are joined with ' | ' and skip empty cells.
        self.assertIn("A | C", content)
        self.assertIn("D", content)
        # Empty row is skipped — only 2 emitted rows.
        self.assertEqual(result["table_row_count"], 2)
        self.assertEqual(result["table_count"], 1)

    def test_table_count_in_return_dict(self) -> None:
        src = self.tmp / "multi_table.docx"
        doc = Document()
        doc.add_paragraph("Heading")
        for _ in range(3):
            t = doc.add_table(rows=2, cols=2)
            t.rows[0].cells[0].text = "x"
            t.rows[0].cells[1].text = "y"
            t.rows[1].cells[0].text = "z"
            t.rows[1].cells[1].text = "w"
        doc.save(str(src))

        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        self.assertIn("table_count", result)
        self.assertIn("table_row_count", result)
        self.assertEqual(result["table_count"], 3)
        self.assertEqual(result["table_row_count"], 6)
        # paragraph_count is now total emitted chunks (1 heading + 6 rows).
        self.assertEqual(result["paragraph_count"], 7)

    def test_pure_table_document_extracted(self) -> None:
        """A .docx with only tables (no paragraphs) must still produce text.

        Reproduces the production bug: meeting minutes that put discussion
        log + action items + next steps in tables previously came out as
        empty_document.
        """
        src = self.tmp / "table_only.docx"
        _write_docx_with_table(
            src,
            [
                ["Decision", "Defer item"],
                ["Owner", "Carol"],
            ],
        )
        result = DocxExtractor().extract(str(src))
        self.assertEqual(result["status"], "success")
        content = Path(result["output_path"]).read_text(encoding="utf-8")
        self.assertIn("Decision | Defer item", content)
        self.assertIn("Owner | Carol", content)


class TestExtractBatch(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_batch_extract_all_docx_in_dir(self) -> None:
        _write_docx(self.tmp / "a.docx", ["Alpha"])
        _write_docx(self.tmp / "b.docx", ["Beta"])
        results = DocxExtractor().extract_batch(str(self.tmp))
        self.assertEqual(len(results), 2)
        self.assertTrue(all(r["status"] == "success" for r in results))
        produced = {Path(r["output_path"]).name for r in results}
        self.assertEqual(produced, {"a.txt", "b.txt"})

    def test_batch_extract_empty_dir_returns_empty_list(self) -> None:
        results = DocxExtractor().extract_batch(str(self.tmp))
        self.assertEqual(results, [])

    def test_batch_extract_custom_output_dir(self) -> None:
        _write_docx(self.tmp / "doc.docx", ["Content"])
        out_dir = self.tmp / "output"
        results = DocxExtractor().extract_batch(str(self.tmp), output_dir=str(out_dir))
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["status"], "success")
        self.assertTrue((out_dir / "doc.txt").is_file())

    def test_batch_ignores_non_docx_files(self) -> None:
        (self.tmp / "readme.txt").write_text("text", encoding="utf-8")
        (self.tmp / "data.json").write_text("{}", encoding="utf-8")
        _write_docx(self.tmp / "real.docx", ["content"])
        results = DocxExtractor().extract_batch(str(self.tmp))
        self.assertEqual(len(results), 1)

    def test_batch_nonexistent_dir_returns_failure_dict(self) -> None:
        results = DocxExtractor().extract_batch(str(self.tmp / "no_such_dir"))
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["status"], "failure")
        self.assertIn("directory_not_found", results[0]["reason"])


class TestWriteError(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_write_error_returns_failure_not_exception(self) -> None:
        src = self.tmp / "doc.docx"
        _write_docx(src, ["Content"])
        dest = self.tmp / "out.txt"
        with patch("spectrum_systems_core.ingestion.docx_extractor.Path.write_text") as mock_write:
            mock_write.side_effect = OSError("disk full")
            result = DocxExtractor().extract(str(src), output_path=str(dest))
        self.assertEqual(result["status"], "failure")
        self.assertIn("write_error", result["reason"])
        self.assertIsNone(result["output_path"])


class TestExtractNeverRaises(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_extract_never_raises(self) -> None:
        src = self.tmp / "broken.docx"
        _write_docx(src, ["something"])
        with patch("spectrum_systems_core.ingestion.docx_extractor.Document") as mock_doc:
            mock_doc.side_effect = RuntimeError("simulated corruption")
            result = DocxExtractor().extract(str(src))
        self.assertIsInstance(result, dict)
        self.assertEqual(result["status"], "failure")
        self.assertIn("docx_parse_error", result["reason"])


class TestCLIExtractDocx(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_cli_extract_single_file(self) -> None:
        src = self.tmp / "meeting.docx"
        _write_docx(src, ["Agenda item one", "Agenda item two"])
        out = io.StringIO()
        rc = extract_docx(path=str(src), out_stream=out)
        self.assertEqual(rc, 0)
        self.assertIn("Extracted:", out.getvalue())
        txt = self.tmp / "meeting.txt"
        self.assertTrue(txt.is_file())

    def test_cli_extract_directory_summary(self) -> None:
        _write_docx(self.tmp / "x.docx", ["X content"])
        _write_docx(self.tmp / "y.docx", ["Y content"])
        out = io.StringIO()
        rc = extract_docx(path=str(self.tmp), out_stream=out)
        self.assertEqual(rc, 0)
        output = out.getvalue()
        self.assertIn("Extracted:", output)
        lines = [l for l in output.splitlines() if "Extracted:" in l]
        self.assertEqual(len(lines), 2)

    def test_cli_exits_1_on_failure(self) -> None:
        out = io.StringIO()
        rc = extract_docx(path="/no/such/path/file.docx", out_stream=out)
        self.assertEqual(rc, 1)
        self.assertIn("error:", out.getvalue())

    def test_cli_empty_dir_exits_0(self) -> None:
        out = io.StringIO()
        rc = extract_docx(path=str(self.tmp), out_stream=out)
        self.assertEqual(rc, 0)
        self.assertIn("No .docx files found", out.getvalue())

    def test_cli_nonexistent_path_exits_1(self) -> None:
        out = io.StringIO()
        rc = extract_docx(path="/completely/nonexistent/dir", out_stream=out)
        self.assertEqual(rc, 1)
        self.assertIn("error:", out.getvalue())
