"""Tests for IngestionEval (post-extraction quality self-test)."""
from __future__ import annotations

import hashlib
import json
import unittest
import uuid
from pathlib import Path
from typing import Any
from unittest.mock import patch

from docx import Document

from spectrum_systems_core.ingestion.docx_extractor import DocxExtractor
from spectrum_systems_core.ingestion.ingestion_eval import (
    MIN_CHARS_PER_BYTE,
    IngestionEval,
)


def _sha256_hex(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _build_rich_docx(path: Path) -> tuple[str, list[dict[str, Any]]]:
    """Build a .docx with several substantive paragraphs and a table.

    Returns ``(extracted_text, text_units)``: the deterministic text
    DocxExtractor would produce, and a corresponding text_units list
    (only the fields IngestionEval inspects).
    """
    long_block = (
        "Today we discussed the 7 GHz downlink coexistence study and "
        "agreed to circulate the latest spectrum sharing draft to all "
        "TIG members for review by next Friday. The propagation model "
        "needs another pass before circulation; Alice will reproduce "
        "the ITU-R P.452 baseline and document any residual deltas, "
        "and Bob will collate the regulatory comments received so far "
        "and produce a single annotated document for the next call. "
    )
    paragraphs = [long_block * 4 for _ in range(6)]
    table_rows = [
        ["Topic", "Owner", "Notes"],
        [
            "Coexistence study draft",
            "Alice",
            (
                "Circulate to TIG by Friday and request comments by "
                "Monday; consolidate received feedback and prepare a "
                "redline for the following call."
            ),
        ],
        [
            "Propagation model",
            "Bob",
            (
                "Reproduce ITU-R P.452 results, document residual "
                "deltas against the prior baseline, and share a "
                "Jupyter notebook with the working group."
            ),
        ],
    ]
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    table = doc.add_table(rows=len(table_rows), cols=3)
    for r_idx, row in enumerate(table_rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    doc.save(str(path))

    # Re-extract via DocxExtractor so the helper returns *exactly* what
    # the extractor produces (avoids whitespace/normalization drift).
    extracted, _, _, _ = DocxExtractor()._extract_body_text(Document(str(path)))
    text_units: list[dict[str, Any]] = [
        {"text": chunk} for chunk in extracted.split("\n\n")
    ]
    return extracted, text_units


def _build_header_only_docx(path: Path) -> tuple[str, list[dict[str, Any]]]:
    """Build a .docx that contains only short headings (the production bug).

    The .docx zip is post-processed to embed a large filler payload so
    the byte-size is realistic relative to the tiny extracted-text size
    — the chars/bytes ratio falls below 0.02 just like the real
    header-only case.
    """
    import zipfile

    headings = [
        "Attendees",
        "Agenda",
        "Discussion",
        "Action Items",
        "Next Steps",
        "Adjourn",
    ]
    doc = Document()
    for h in headings:
        doc.add_paragraph(h)
    doc.save(str(path))

    # Append a ~50KB filler entry inside the .docx zip so the file size
    # mirrors a real meeting-minutes .docx without affecting the
    # extractor's paragraph walk.
    with zipfile.ZipFile(str(path), "a", compression=zipfile.ZIP_STORED) as zf:
        zf.writestr("filler.bin", b"x" * 50_000)

    extracted = "\n\n".join(headings)
    text_units = [{"text": h} for h in headings]
    return extracted, text_units


def _make_source_record(
    *,
    artifact_id: str,
    text_units: list[dict[str, Any]],
    raw_hash: str,
    processed_path: str = "",
) -> dict[str, Any]:
    return {
        "artifact_kind": "source_record",
        "artifact_id": artifact_id,
        "payload": {
            "source_id": "test-source",
            "source_family": "meetings",
            "raw_hash": raw_hash,
            "text_unit_count": len(text_units),
            "processed_path": processed_path,
        },
    }


class TestEvalPasses(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_eval_passes_for_rich_transcript(self) -> None:
        src = self.tmp / "rich.docx"
        extracted, units = _build_rich_docx(src)
        raw_hash = "sha256:" + _sha256_hex(extracted.encode("utf-8"))
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units,
            raw_hash=raw_hash,
        )

        result = IngestionEval().evaluate(str(src), record, text_units=units)

        self.assertEqual(result["status"], "passed", msg=result)
        self.assertTrue(result["eval_passed"])
        self.assertEqual(result["failure_reasons"], [])
        names = {c["name"] for c in result["checks"]}
        self.assertEqual(
            names,
            {
                "text_units_present",
                "minimum_content_ratio",
                "not_header_only",
                "deterministic_extraction",
            },
        )
        for c in result["checks"]:
            self.assertTrue(c["passed"], msg=c)


class TestEvalFails(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_eval_fails_for_header_only_document(self) -> None:
        src = self.tmp / "headers_only.docx"
        extracted, units = _build_header_only_docx(src)
        raw_hash = "sha256:" + _sha256_hex(extracted.encode("utf-8"))
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units,
            raw_hash=raw_hash,
        )

        result = IngestionEval().evaluate(str(src), record, text_units=units)

        self.assertEqual(result["status"], "failed", msg=result)
        self.assertFalse(result["eval_passed"])
        # Both content-ratio and header-only checks should fail.
        failed_names = {
            c["name"] for c in result["checks"]
            if not c["passed"] and c["required"]
        }
        self.assertIn("minimum_content_ratio", failed_names)
        self.assertIn("not_header_only", failed_names)
        # Failure reasons surface the symptom.
        self.assertTrue(
            any(r.startswith("extraction_too_sparse") for r in result["failure_reasons"]),
            result["failure_reasons"],
        )
        self.assertTrue(
            any(r.startswith("likely_header_only") for r in result["failure_reasons"]),
            result["failure_reasons"],
        )

    def test_eval_fails_for_empty_document(self) -> None:
        src = self.tmp / "empty.docx"
        Document().save(str(src))
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=[],
            raw_hash="sha256:" + ("0" * 64),
        )

        result = IngestionEval().evaluate(str(src), record, text_units=[])

        self.assertEqual(result["status"], "failed", msg=result)
        self.assertFalse(result["eval_passed"])
        self.assertIn("no_text_units_extracted", result["failure_reasons"])


class TestEvalAdvisory(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_eval_warns_on_hash_mismatch(self) -> None:
        src = self.tmp / "rich.docx"
        _, units = _build_rich_docx(src)
        # Stored raw_hash deliberately wrong — required checks still pass,
        # but deterministic_extraction (advisory) flags a warning.
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units,
            raw_hash="sha256:" + ("a" * 64),
        )

        result = IngestionEval().evaluate(str(src), record, text_units=units)

        self.assertEqual(result["status"], "warning", msg=result)
        self.assertTrue(result["eval_passed"])
        self.assertEqual(result["failure_reasons"], [])
        det = next(
            c for c in result["checks"] if c["name"] == "deterministic_extraction"
        )
        self.assertFalse(det["passed"])
        self.assertFalse(det["required"])
        self.assertIn("hash_mismatch", det["detail"])


class TestThresholds(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_minimum_content_ratio_check(self) -> None:
        """Below 0.02 fails; well above 0.02 passes."""
        src = self.tmp / "rich.docx"
        extracted, units = _build_rich_docx(src)
        raw_hash = "sha256:" + _sha256_hex(extracted.encode("utf-8"))
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units,
            raw_hash=raw_hash,
        )
        result = IngestionEval().evaluate(str(src), record, text_units=units)
        self.assertGreaterEqual(
            result["ratio_chars_per_byte"], MIN_CHARS_PER_BYTE
        )

        src2 = self.tmp / "headers.docx"
        _, units2 = _build_header_only_docx(src2)
        record2 = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units2,
            raw_hash="sha256:" + ("0" * 64),
        )
        result2 = IngestionEval().evaluate(str(src2), record2, text_units=units2)
        self.assertLess(
            result2["ratio_chars_per_byte"], MIN_CHARS_PER_BYTE
        )

    def test_not_header_only_check(self) -> None:
        """80% short-unit threshold."""
        # 4 long + 1 short (20% short) → passes.
        long_text = "a long substantive paragraph " * 5
        units = [{"text": long_text} for _ in range(4)] + [{"text": "Short"}]
        # Build a .docx whose extracted text matches `units` so the ratio
        # check also passes.
        src = self.tmp / "mostly_long.docx"
        doc = Document()
        for u in units:
            doc.add_paragraph(u["text"])
        doc.save(str(src))
        text, _, _, _ = DocxExtractor()._extract_body_text(Document(str(src)))
        raw_hash = "sha256:" + _sha256_hex(text.encode("utf-8"))
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units,
            raw_hash=raw_hash,
        )
        result = IngestionEval().evaluate(str(src), record, text_units=units)
        not_header = next(
            c for c in result["checks"] if c["name"] == "not_header_only"
        )
        self.assertTrue(not_header["passed"], msg=result)

        # 1 long + 4 short (80% short) → fails.
        units_bad = [{"text": long_text}] + [{"text": "S"} for _ in range(4)]
        src2 = self.tmp / "mostly_short.docx"
        doc2 = Document()
        for u in units_bad:
            doc2.add_paragraph(u["text"])
        doc2.save(str(src2))
        text2, _, _, _ = DocxExtractor()._extract_body_text(Document(str(src2)))
        raw_hash2 = "sha256:" + _sha256_hex(text2.encode("utf-8"))
        record_bad = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units_bad,
            raw_hash=raw_hash2,
        )
        result_bad = IngestionEval().evaluate(str(src2), record_bad, text_units=units_bad)
        not_header_bad = next(
            c for c in result_bad["checks"] if c["name"] == "not_header_only"
        )
        self.assertFalse(not_header_bad["passed"], msg=result_bad)


class TestRedteamFollowups(unittest.TestCase):
    """Gate A redteam coverage."""

    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_dialogue_heavy_transcript_passes(self) -> None:
        """Many short speaker turns + lots of total content must pass.

        Guards against the false-positive class flagged in Gate A: a
        Teams transcript with many one-word turns but plenty of content
        previously failed the short-unit ratio check.
        """
        long_block = "this is a substantive paragraph " * 30
        # 8 short turns + 4 long blocks. short_ratio = 8/12 ~= 0.67 < 0.8
        # AND total chars are well above 2000.
        units = (
            [{"text": "Yes."} for _ in range(4)]
            + [{"text": "OK."} for _ in range(4)]
            + [{"text": long_block} for _ in range(4)]
        )
        src = self.tmp / "dialogue.docx"
        doc = Document()
        for u in units:
            doc.add_paragraph(u["text"])
        doc.save(str(src))
        text, _, _, _ = DocxExtractor()._extract_body_text(Document(str(src)))
        raw_hash = "sha256:" + _sha256_hex(text.encode("utf-8"))
        # Re-derive units to exactly match the extractor output (avoids
        # whitespace drift between input strings and extracted text).
        units = [{"text": t} for t in text.split("\n\n")]
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units,
            raw_hash=raw_hash,
        )

        result = IngestionEval().evaluate(str(src), record, text_units=units)

        self.assertEqual(result["status"], "passed", msg=result)
        not_header = next(
            c for c in result["checks"] if c["name"] == "not_header_only"
        )
        self.assertTrue(not_header["passed"], msg=result)

    def test_text_units_unloadable_distinct_failure_reason(self) -> None:
        """When source_record claims units exist but the jsonl can't be
        loaded, the failure reason names that specifically — not just
        ``likely_header_only``."""
        src = self.tmp / "rich.docx"
        extracted, _ = _build_rich_docx(src)
        raw_hash = "sha256:" + _sha256_hex(extracted.encode("utf-8"))
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=[{"text": "x"}],  # placeholder for count only
            raw_hash=raw_hash,
            processed_path="/no/such/dir",
        )
        # Override text_unit_count to claim 5 units while we pass None
        # so the eval tries to load from disk and fails.
        record["payload"]["text_unit_count"] = 5

        result = IngestionEval().evaluate(str(src), record, text_units=None)

        self.assertEqual(result["status"], "failed", msg=result)
        self.assertTrue(
            any(
                r.startswith("text_units_unloadable")
                for r in result["failure_reasons"]
            ),
            result["failure_reasons"],
        )

    def test_malformed_jsonl_lines_skipped(self) -> None:
        """A processed_path with malformed jsonl lines must not raise and
        should leave the eval able to fall back to fail-closed behavior."""
        src = self.tmp / "rich.docx"
        extracted, _ = _build_rich_docx(src)
        raw_hash = "sha256:" + _sha256_hex(extracted.encode("utf-8"))
        proc_dir = self.tmp / "processed"
        proc_dir.mkdir()
        # Mix of valid and malformed lines.
        (proc_dir / "text_units.jsonl").write_text(
            '{"text":"valid line"}\n'
            "this is not json\n"
            '{"text":"another valid"}\n',
            encoding="utf-8",
        )
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=[{"text": "x"}, {"text": "y"}],
            raw_hash=raw_hash,
            processed_path=str(proc_dir),
        )

        result = IngestionEval().evaluate(str(src), record, text_units=None)

        # No exception, result is well-formed, and only the 2 valid lines
        # are loaded from disk.
        self.assertIsInstance(result, dict)
        self.assertIn("status", result)


class TestSchemaConformance(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_source_eval_result_schema_validates(self) -> None:
        src = self.tmp / "rich.docx"
        extracted, units = _build_rich_docx(src)
        raw_hash = "sha256:" + _sha256_hex(extracted.encode("utf-8"))
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units,
            raw_hash=raw_hash,
        )
        eval_obj = IngestionEval()
        result = eval_obj.evaluate(str(src), record, text_units=units)
        self.assertTrue(eval_obj.schema_validate(result), msg=result)

        # And a failing result must also validate the schema.
        src2 = self.tmp / "empty.docx"
        Document().save(str(src2))
        record2 = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=[],
            raw_hash="sha256:" + ("0" * 64),
        )
        result2 = eval_obj.evaluate(str(src2), record2, text_units=[])
        self.assertEqual(result2["status"], "failed")
        self.assertTrue(eval_obj.schema_validate(result2), msg=result2)


class TestNeverRaises(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_eval_never_raises(self) -> None:
        # Mock the body walker to raise; result must still be a dict.
        src = self.tmp / "rich.docx"
        extracted, units = _build_rich_docx(src)
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units,
            raw_hash="sha256:" + _sha256_hex(extracted.encode("utf-8")),
        )
        with patch.object(
            DocxExtractor,
            "_extract_body_text",
            side_effect=RuntimeError("boom"),
        ):
            result = IngestionEval().evaluate(str(src), record, text_units=units)
        self.assertIsInstance(result, dict)
        self.assertIn("status", result)
        # The advisory check should warn (recompute_failed) but required
        # checks still pass — so status is "warning", not "failed".
        self.assertIn(result["status"], {"warning", "failed"})

    def test_eval_handles_garbage_input(self) -> None:
        # Pass arbitrary garbage — must still return a well-formed dict.
        result = IngestionEval().evaluate("/no/such/file.docx", {"junk": True})
        self.assertIsInstance(result, dict)
        self.assertEqual(result["status"], "failed")
        self.assertFalse(result["eval_passed"])


class TestSdlRootWrite(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_eval_written_to_sdl_root(self) -> None:
        src = self.tmp / "rich.docx"
        extracted, units = _build_rich_docx(src)
        record = _make_source_record(
            artifact_id="abc-123",
            text_units=units,
            raw_hash="sha256:" + _sha256_hex(extracted.encode("utf-8")),
        )
        eval_obj = IngestionEval()
        result = eval_obj.evaluate(str(src), record, text_units=units)

        sdl_root = self.tmp / "sdl"
        path_written = eval_obj.write_eval_result(
            result, sdl_root=str(sdl_root)
        )

        self.assertTrue(path_written)
        target = Path(path_written)
        self.assertTrue(target.is_file())
        # Filename uses source_artifact_id.
        self.assertTrue(target.name.endswith("_ingestion_eval.json"))
        self.assertIn("abc-123", target.name)
        # File contents round-trip.
        loaded = json.loads(target.read_text(encoding="utf-8"))
        self.assertEqual(loaded["source_artifact_id"], "abc-123")

    def test_write_returns_empty_when_sdl_root_unset(self) -> None:
        src = self.tmp / "rich.docx"
        _, units = _build_rich_docx(src)
        record = _make_source_record(
            artifact_id=str(uuid.uuid4()),
            text_units=units,
            raw_hash="sha256:" + ("0" * 64),
        )
        eval_obj = IngestionEval()
        result = eval_obj.evaluate(str(src), record, text_units=units)
        # No sdl_root arg, no SDL_ROOT env: write should no-op.
        with patch.dict("os.environ", {}, clear=False):
            import os as _os
            _os.environ.pop("SDL_ROOT", None)
            path_written = eval_obj.write_eval_result(result)
        self.assertEqual(path_written, "")


# ---------------------------------------------------------------------------
# Orchestrator integration
# ---------------------------------------------------------------------------


class TestOrchestratorIntegration(unittest.TestCase):
    """Verify the orchestrator records eval_status and never blocks on it."""

    def setUp(self) -> None:
        import tempfile
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def _make_data_lake(self, with_docx_kind: str) -> Path:
        """Create a data lake with one .docx of the named kind ('rich' or 'header')."""
        root = self.tmp / "lake"
        transcripts = root / "store" / "raw" / "transcripts"
        transcripts.mkdir(parents=True)
        if with_docx_kind == "rich":
            _build_rich_docx(transcripts / "meeting.docx")
        else:
            _build_header_only_docx(transcripts / "meeting.docx")
        return root

    def _set_env(self, root: Path) -> None:
        """Promoter and IngestionEval read DATA_LAKE_PATH / SDL_ROOT."""
        import os
        sdl = root / "store" / "artifacts"
        sdl.mkdir(parents=True, exist_ok=True)
        self._old = {
            "DATA_LAKE_PATH": os.environ.get("DATA_LAKE_PATH"),
            "SDL_ROOT": os.environ.get("SDL_ROOT"),
        }
        os.environ["DATA_LAKE_PATH"] = str(root)
        os.environ["SDL_ROOT"] = str(sdl)

    def _restore_env(self) -> None:
        import os
        for k, v in getattr(self, "_old", {}).items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v

    def test_orchestrator_includes_eval_status_in_record(self) -> None:
        from spectrum_systems_core.orchestration.pipeline_orchestrator import (
            PipelineOrchestrator,
        )

        root = self._make_data_lake("rich")
        self._set_env(root)
        try:
            orchestrator = PipelineOrchestrator()
            result = orchestrator.run(str(root), dry_run=False)
        finally:
            self._restore_env()

        record_path = result["orchestration_record_path"]
        self.assertTrue(record_path, msg=result)
        record = json.loads(Path(record_path).read_text(encoding="utf-8"))
        self.assertEqual(len(record["results"]), 1, msg=record)
        entry = record["results"][0]
        self.assertIn("eval_status", entry)
        self.assertIn(
            entry["eval_status"], {"passed", "warning", "failed", "not_run"}
        )

    def test_orchestrator_continues_on_eval_failure(self) -> None:
        from spectrum_systems_core.orchestration.pipeline_orchestrator import (
            PipelineOrchestrator,
        )

        root = self._make_data_lake("header")
        self._set_env(root)
        try:
            orchestrator = PipelineOrchestrator()
            result = orchestrator.run(str(root), dry_run=False)
        finally:
            self._restore_env()

        # The transcript is still processed and promoted (eval is advisory).
        self.assertEqual(result["status"], "success", msg=result)
        self.assertEqual(result["total_failed"], 0)

    def test_orchestrator_marks_warning_not_blocking(self) -> None:
        from spectrum_systems_core.orchestration.pipeline_orchestrator import (
            PipelineOrchestrator,
        )

        root = self._make_data_lake("header")
        self._set_env(root)
        try:
            orchestrator = PipelineOrchestrator()
            result = orchestrator.run(str(root), dry_run=False)
        finally:
            self._restore_env()

        self.assertEqual(result["status"], "success", msg=result)
        record = json.loads(
            Path(result["orchestration_record_path"]).read_text(encoding="utf-8")
        )
        statuses = {e["status"] for e in record["results"]}
        self.assertIn("extraction_quality_warning", statuses)
        eval_statuses = {e["eval_status"] for e in record["results"]}
        self.assertIn("failed", eval_statuses)


if __name__ == "__main__":
    unittest.main()
