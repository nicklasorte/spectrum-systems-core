"""Unit tests for ``scripts/create_opus_reference_baselines.py``.

Every LLM call is a stubbed in-process client (``(*, system, user) ->
str``) — the SAME structural seam ``workflows/llm_client.py`` defines.
No network, no ``ANTHROPIC_API_KEY``. Each gate is tested on both its
happy path and its rejection path.
"""
from __future__ import annotations

import json
import subprocess
import sys
import uuid
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = REPO_ROOT / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import create_opus_reference_baselines as cob  # noqa: E402

from spectrum_systems_core.promotion.gate import (  # noqa: E402
    GROUNDING_BINDING_SCHEMA_VERSION,
)
from tests.integration.fixtures import make_source_record  # noqa: E402

MODEL = "claude-opus-4-6"
TRANSCRIPT_STEM = "sample-transcript-20251218"
SOURCE_ID = "sample-transcript-20251218"  # _slugify(stem) == stem here

_VALID_RESPONSE = json.dumps(
    {
        "decisions": [
            "The group approved the 7 GHz downlink threshold.",
            "The group deferred the aggregate methodology.",
        ],
        "action_items": ["DoD will submit revised ERP values."],
        "open_questions": ["What is the coordination distance?"],
        "commitments": [
            {
                "commitment_id": "commit-1",
                "owner": "DoD Rep",
                "commitment_text": "DoD will submit revised ERP values.",
                "due": None,
                "source_speaker": "DoD Rep",
            }
        ],
        "risks": [
            {
                "risk_id": "risk-1",
                "risk_text": "Concern about the methodology.",
                "raised_by": "DoD Rep",
                "severity": None,
                "mitigation_mentioned": None,
            }
        ],
        "cross_references": [],
        "attendees": [
            {
                "name": "Chair Smith",
                "agency": "FCC",
                "role": "Chair",
                "present": True,
            }
        ],
        "topics": [],
        "regulatory_references": [],
        "technical_parameters": [
            {
                "param_id": "param-1",
                "parameter_name": "7 GHz downlink threshold",
                "value": "minus 47 dBm per megahertz",
                "unit": "dBm/MHz",
                "context": "approved threshold",
                "speaker": "NTIA Lead",
            }
        ],
        "named_artifacts": [],
        "scheduled_events": [],
    }
)


class SpyClient:
    """Records call count and the exact ``system`` prompt it received."""

    def __init__(self, response: str):
        self.response = response
        self.calls = 0
        self.last_system: str | None = None
        self.last_user: str | None = None

    def __call__(self, *, system: str, user: str) -> str:
        self.calls += 1
        self.last_system = system
        self.last_user = user
        return self.response


class SequenceClient:
    """Returns a different response per call (1st, 2nd, ...).

    Drives the Fix 3 retry path: the first call returns a malformed
    blob, the second returns the simplified-retry reply. Records every
    ``user`` message so a test can prove the retry sent ONLY the first
    200 chars of the failed response (never the whole broken blob).
    """

    def __init__(self, responses: list[str]):
        self._responses = list(responses)
        self.calls = 0
        self.users: list[str] = []
        self.systems: list[str] = []

    def __call__(self, *, system: str, user: str) -> str:
        self.systems.append(system)
        self.users.append(user)
        idx = min(self.calls, len(self._responses) - 1)
        self.calls += 1
        return self._responses[idx]


def _make_docx(path: Path) -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("7 GHz Downlink TIG Kickoff Meeting")
    doc.add_paragraph("Meeting Date: 2025-12-18")
    doc.add_paragraph(
        "The group approved the 7 GHz downlink threshold of minus 47 "
        "dBm per megahertz."
    )
    doc.add_paragraph("DoD will submit revised ERP values.")
    doc.save(str(path))


def _seed(tmp_path: Path, *, with_source_record: bool = True) -> Path:
    dl = tmp_path / "data-lake"
    _make_docx(
        dl / "store" / "raw" / "transcripts" / f"{TRANSCRIPT_STEM}.docx"
    )
    if with_source_record:
        proc = dl / "store" / "processed" / "meetings" / SOURCE_ID
        proc.mkdir(parents=True)
        (proc / "source_record.json").write_text(
            json.dumps(make_source_record(SOURCE_ID, str(uuid.uuid4())))
        )
    return dl


def _out_path(dl: Path) -> Path:
    return (
        dl / "store" / "processed" / "meetings" / SOURCE_ID
        / "reference_baselines" / "opus_reference_minutes.jsonl"
    )


# --------------------------------------------------------------------------
# 1. Dry run produces no files
# --------------------------------------------------------------------------
def test_dry_run_produces_no_files(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    spy = SpyClient(_VALID_RESPONSE)
    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=True,
        skip_existing=True,
        model=MODEL,
        client=spy,
    )
    assert result["status"] == "success"
    assert not _out_path(dl).exists()
    t = result["per_transcript"][0]
    assert t["status"] == "dry_run"
    assert t["total"] > 0


# --------------------------------------------------------------------------
# 2. Skip-existing skips a source that already has the JSONL
# --------------------------------------------------------------------------
def test_skip_existing_skips_present_jsonl(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    out = _out_path(dl)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text('{"pre":"existing"}\n', encoding="utf-8")

    spy = SpyClient(_VALID_RESPONSE)
    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=spy,
    )
    assert spy.calls == 0, "skip-existing must not even call the model"
    assert result["per_transcript"][0]["status"] == "skipped"
    assert out.read_text(encoding="utf-8") == '{"pre":"existing"}\n'


def test_no_skip_existing_regenerates(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    out = _out_path(dl)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text('{"pre":"existing"}\n', encoding="utf-8")

    spy = SpyClient(_VALID_RESPONSE)
    cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=False,
        model=MODEL,
        client=spy,
    )
    assert spy.calls == 1
    assert '"pre":"existing"' not in out.read_text(encoding="utf-8")


# --------------------------------------------------------------------------
# 3. Valid transcript -> JSONL with correct schema fields
#    7. model_id matches --model    8. authorship flags
# --------------------------------------------------------------------------
def test_valid_transcript_writes_correct_fields(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(_VALID_RESPONSE),
    )
    lines = [
        json.loads(ln)
        for ln in _out_path(dl).read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]
    # 2 decisions + 1 action + 1 question + 1 commitment + 1 risk
    # + 1 attendee + 1 technical_parameter = 8
    assert len(lines) == 8

    required = {
        "pair_id", "source_id", "source_artifact_id", "extraction_type",
        "ground_truth_text", "item_data", "human_authored",
        "model_authored", "model_id", "verified", "status",
        "provenance", "schema_version", "meeting_date", "created_at",
        "chunking_strategy_version",
    }
    for r in lines:
        assert required <= set(r), f"missing fields: {required - set(r)}"
        assert r["model_id"] == MODEL                       # (7)
        assert r["human_authored"] is False                 # (8)
        assert r["model_authored"] is True                  # (8)
        assert r["verified"] is False
        assert r["status"] == "reference_only"
        assert r["provenance"] == {
            "produced_by": "opus_reference_baseline_workflow"
        }
        # Read from the canonical source so the test does not drift
        # when the gate's binding schema_version bumps. A string
        # literal here is the bug-class that produced the
        # schema_version_mixed halt: the writer stayed at "1.0.0" while
        # the rest of the pipeline advanced to "1.4.0".
        assert r["schema_version"] == GROUNDING_BINDING_SCHEMA_VERSION
        assert r["source_id"] == SOURCE_ID
        assert r["meeting_date"] == "2025-12-18"  # from docx header
        assert r["ground_truth_text"].strip()
        assert r["created_at"].endswith("+00:00")
        # Phase 2.B: default `CHUNK_OVERLAP_TURNS` unset/0 yields the
        # no-suffix strategy token, matching the comparator's default
        # for any baseline row that omits the field.
        assert r["chunking_strategy_version"] == "speaker_turn_v1"

    by_type: dict[str, list] = {}
    for r in lines:
        by_type.setdefault(r["extraction_type"], []).append(r)
    # String array: ground_truth_text IS the string.
    assert sorted(x["ground_truth_text"] for x in by_type["decisions"]) == [
        "The group approved the 7 GHz downlink threshold.",
        "The group deferred the aggregate methodology.",
    ]
    # Structured items: ground_truth_text is the first matching field
    # from the priority list in ``extract_ground_truth_text`` —
    # ``risk_text`` for a risk, ``name`` for an attendee, and
    # ``parameter_name`` (which precedes ``value`` semantics, and is the
    # only ``technical_parameters`` field in the priority list) for a
    # technical parameter.
    assert by_type["risks"][0]["ground_truth_text"] == (
        "Concern about the methodology."
    )
    assert by_type["attendees"][0]["ground_truth_text"] == "Chair Smith"
    assert by_type["technical_parameters"][0]["ground_truth_text"] == (
        "7 GHz downlink threshold"
    )
    # item_data is the full structured object, verbatim.
    assert by_type["risks"][0]["item_data"]["risk_id"] == "risk-1"
    # A string item's item_data is wrapped as {"text": <string>} so the
    # original value is always recoverable from a JSON object.
    assert by_type["decisions"][0]["item_data"] == {
        "text": by_type["decisions"][0]["ground_truth_text"]
    }


# --------------------------------------------------------------------------
# 3b. Phase 2.B: CHUNK_OVERLAP_TURNS env var stamps the matching
#     chunking_strategy_version on every row, and the summary dict
#     surfaces both the raw env value and the resolved version.
# --------------------------------------------------------------------------
def test_chunk_overlap_turns_env_stamps_version_on_rows(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("CHUNK_OVERLAP_TURNS", "2")
    dl = _seed(tmp_path)
    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(_VALID_RESPONSE),
    )
    assert result["chunk_overlap_turns"] == "2"
    assert result["chunking_strategy_version"] == "speaker_turn_v1_overlap2"

    lines = [
        json.loads(ln)
        for ln in _out_path(dl).read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]
    assert lines, "expected at least one record"
    for r in lines:
        assert r["chunking_strategy_version"] == "speaker_turn_v1_overlap2"


def test_chunk_overlap_turns_default_yields_no_suffix(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Default unset/0 must produce the no-suffix `speaker_turn_v1`
    token, byte-compatible with pre-Phase-2.B Opus baselines on disk
    that omit the field (the comparator defaults a missing value to
    the same string)."""
    monkeypatch.delenv("CHUNK_OVERLAP_TURNS", raising=False)
    dl = _seed(tmp_path)
    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(_VALID_RESPONSE),
    )
    assert result["chunk_overlap_turns"] == "0"
    assert result["chunking_strategy_version"] == "speaker_turn_v1"
    for ln in _out_path(dl).read_text(encoding="utf-8").splitlines():
        if ln.strip():
            assert (
                json.loads(ln)["chunking_strategy_version"]
                == "speaker_turn_v1"
            )


# --------------------------------------------------------------------------
# 3c. Opus baseline schema_version mismatch regression. Reproduces the
#     failure mode where the writer hard-coded "1.0.0" while the
#     pipeline's grounding-binding schema_version had advanced to
#     "1.4.0", producing the `schema_version_mixed` halt in the
#     comparison engine. The test asserts the writer reads from the
#     canonical source (`gate.GROUNDING_BINDING_SCHEMA_VERSION`), so the
#     next schema bump flows through automatically without re-touching
#     this script.
# --------------------------------------------------------------------------
def test_opus_baseline_schema_version_matches_canonical_source_no_string_literal(
    tmp_path: Path,
) -> None:
    """Pin the Opus reference baseline schema_version to the canonical
    `GROUNDING_BINDING_SCHEMA_VERSION` constant.

    Two assertions:

    1. Every written row stamps the canonical version (positive
       behaviour).
    2. No row stamps the legacy "1.0.0" literal (regression guard
       against the schema_version_mismatch foot-gun: a hard-coded
       string left behind by a schema bump silently mis-tags every
       baseline row and the comparator's `_baseline_at_version_exists`
       check then refuses every comparison with `schema_version_mixed`).
    """
    dl = _seed(tmp_path)
    cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(_VALID_RESPONSE),
    )
    lines = [
        json.loads(ln)
        for ln in _out_path(dl).read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]
    assert lines, "expected at least one record"
    for r in lines:
        # Positive: stamped value equals the canonical constant.
        assert r["schema_version"] == GROUNDING_BINDING_SCHEMA_VERSION, (
            f"row stamps {r['schema_version']!r}; expected canonical "
            f"{GROUNDING_BINDING_SCHEMA_VERSION!r}"
        )
    # The canonical constant itself must not have silently reverted to
    # the pre-Phase-1 legacy "1.0.0" — if it does, every comparison
    # against a 1.4.0+ Haiku artifact halts on schema_version_mixed.
    assert GROUNDING_BINDING_SCHEMA_VERSION != "1.0.0", (
        "GROUNDING_BINDING_SCHEMA_VERSION reverted to legacy 1.0.0; "
        "this re-introduces the schema_version_mismatch bug class."
    )


# --------------------------------------------------------------------------
# 4. Missing source_record.json -> deterministic Stage-1 self-heal
#    (LLM-free) produces it, then the baseline is written. The unstated
#    "transcript must be pre-ingested" precondition no longer halts the
#    run on every invocation.
# --------------------------------------------------------------------------
def test_missing_source_record_auto_ingests(tmp_path: Path) -> None:
    dl = _seed(tmp_path, with_source_record=False)
    sr_path = (
        dl / "store" / "processed" / "meetings" / SOURCE_ID
        / "source_record.json"
    )
    assert not sr_path.exists(), "precondition: no source_record yet"

    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(_VALID_RESPONSE),
    )

    assert result["status"] == "success"
    # The reused SourceLoader (the real writer) produced the record.
    assert sr_path.is_file(), "Stage-1 ingestion must create the record"
    record = json.loads(sr_path.read_text(encoding="utf-8"))
    aid = record["artifact_id"]
    assert uuid.UUID(aid)  # valid UUID, as the pipeline writes it

    out = _out_path(dl)
    assert out.is_file(), "baseline must be written after self-heal"
    lines = [
        json.loads(ln)
        for ln in out.read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]
    assert lines, "non-empty baseline expected"
    for r in lines:
        assert r["source_artifact_id"] == aid


def test_present_but_invalid_source_record_still_halts(
    tmp_path: Path,
) -> None:
    """Fail-closed NOT weakened: a *present* source_record with no
    artifact_id still halts ``invalid_source_record`` — the self-heal
    only runs when the record is ABSENT, never over a corrupt one."""
    dl = _seed(tmp_path, with_source_record=False)
    proc = dl / "store" / "processed" / "meetings" / SOURCE_ID
    proc.mkdir(parents=True)
    (proc / "source_record.json").write_text(
        json.dumps({"source_id": SOURCE_ID}), encoding="utf-8"
    )
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=SpyClient(_VALID_RESPONSE),
        )
    assert ei.value.reason == "invalid_source_record"
    assert not _out_path(dl).exists()


def _seed_custom_record(tmp_path: Path, record: dict) -> Path:
    """Seed the data-lake with a hand-written source_record.json.

    The minimal-contract gate cases below assert that ONLY a
    valid-UUID ``artifact_id`` is required, so they intentionally do
    NOT route through ``make_source_record`` — they pin the exact
    record shapes the contract must accept or reject.
    """
    dl = _seed(tmp_path, with_source_record=False)
    proc = dl / "store" / "processed" / "meetings" / SOURCE_ID
    proc.mkdir(parents=True)
    (proc / "source_record.json").write_text(
        json.dumps(record), encoding="utf-8"
    )
    return dl


def test_source_record_only_artifact_id_passes(tmp_path: Path) -> None:
    """Gate: a record carrying ONLY a valid-UUID artifact_id passes."""
    aid = str(uuid.uuid4())
    dl = _seed_custom_record(tmp_path, {"artifact_id": aid})
    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(_VALID_RESPONSE),
    )
    assert result["status"] == "success"
    out = _out_path(dl)
    assert out.is_file()
    for ln in out.read_text(encoding="utf-8").splitlines():
        if ln.strip():
            assert json.loads(ln)["source_artifact_id"] == aid


def test_source_record_artifact_id_plus_source_id_passes(
    tmp_path: Path,
) -> None:
    """Gate: artifact_id + a redundant top-level source_id still passes."""
    aid = str(uuid.uuid4())
    dl = _seed_custom_record(
        tmp_path, {"artifact_id": aid, "source_id": SOURCE_ID}
    )
    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(_VALID_RESPONSE),
    )
    assert result["status"] == "success"
    assert _out_path(dl).is_file()


def test_source_record_missing_artifact_id_halts(tmp_path: Path) -> None:
    """Gate: a record with no artifact_id halts invalid_source_record."""
    dl = _seed_custom_record(tmp_path, {"source_id": SOURCE_ID})
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=SpyClient(_VALID_RESPONSE),
        )
    assert ei.value.reason == "invalid_source_record"
    assert not _out_path(dl).exists()


def test_source_record_non_uuid_artifact_id_halts(tmp_path: Path) -> None:
    """A present but non-UUID artifact_id is still a fail-closed halt."""
    dl = _seed_custom_record(tmp_path, {"artifact_id": "x"})
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=SpyClient(_VALID_RESPONSE),
        )
    assert ei.value.reason == "invalid_source_record"
    assert not _out_path(dl).exists()


# --------------------------------------------------------------------------
# 5. Missing prompt -> halt with missing_extraction_prompt BEFORE the LLM
# --------------------------------------------------------------------------
def test_missing_prompt_halts_before_llm(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    dl = _seed(tmp_path)
    spy = SpyClient(_VALID_RESPONSE)
    monkeypatch.setattr(
        cob, "_PROMPT_PATH", tmp_path / "does-not-exist.md"
    )
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=spy,
        )
    assert ei.value.reason == "missing_extraction_prompt"
    assert spy.calls == 0, "must halt BEFORE any model call"
    assert not _out_path(dl).exists()


def test_empty_prompt_halts(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    dl = _seed(tmp_path)
    empty = tmp_path / "empty.md"
    empty.write_text("   \n", encoding="utf-8")
    spy = SpyClient(_VALID_RESPONSE)
    monkeypatch.setattr(cob, "_PROMPT_PATH", empty)
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=spy,
        )
    assert ei.value.reason == "missing_extraction_prompt"
    assert spy.calls == 0


# --------------------------------------------------------------------------
# 6. Malformed LLM response -> halt, no write
# --------------------------------------------------------------------------
# Only RESPONSE-level malformations halt now. A structured item with
# no schema-named text field is NOT a halt — it is tolerantly read by
# ``extract_ground_truth_text`` (the canonical prompt explicitly allows
# a structured object for every type), so the old
# ``{"risks": [{"risk_id": "r1"}]}`` "missing required text" halt case
# was intentionally removed; that shape is now covered by
# ``test_extract_ground_truth_text_*`` as a recoverable item.
@pytest.mark.parametrize(
    "bad",
    [
        "this is not json at all",
        "[1, 2, 3]",                       # JSON but not an object
        '{"decisions": "oops not a list"}',  # content key not a list
        "",                                 # empty text
    ],
)
def test_malformed_llm_response_halts_no_write(
    tmp_path: Path, bad: str
) -> None:
    dl = _seed(tmp_path)
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=SpyClient(bad),
        )
    assert ei.value.reason == "malformed_llm_response"
    assert not _out_path(dl).exists(), "no partial JSONL on malformed"


# --------------------------------------------------------------------------
# Structured-object item handling: extract_ground_truth_text NEVER halts
# (the canonical prompt allows a structured object for every type).
# --------------------------------------------------------------------------
def test_extract_ground_truth_text_decisions_string() -> None:
    """decisions item as a plain string -> the string verbatim."""
    assert (
        cob.extract_ground_truth_text("NTIA approved X.", "decisions")
        == "NTIA approved X."
    )


def test_extract_ground_truth_text_decisions_object() -> None:
    """The exact Opus shape that USED to halt now resolves on ``text``."""
    item = {
        "text": "The TIG approved the 7 GHz downlink threshold.",
        "verb": "decided",
        "stakeholders": ["DoD"],
        "confidence": 0.85,
        "rationale": None,
    }
    assert cob.extract_ground_truth_text(item, "decisions") == (
        "The TIG approved the 7 GHz downlink threshold."
    )


def test_extract_ground_truth_text_action_item_object() -> None:
    """action_items dict with ``text`` + ``owner`` -> the text field."""
    item = {"text": "NTIA to circulate the methodology.", "owner": "NTIA"}
    assert cob.extract_ground_truth_text(item, "action_items") == (
        "NTIA to circulate the methodology."
    )


def test_extract_ground_truth_text_open_question_object() -> None:
    """open_questions dict with ``question_text`` -> that field."""
    item = {"question_text": "What is the coordination distance?"}
    assert cob.extract_ground_truth_text(item, "open_questions") == (
        "What is the coordination distance?"
    )


def test_extract_ground_truth_text_risk_object() -> None:
    """risks dict with ``risk_text`` -> that field (not a halt)."""
    item = {"risk_id": "risk-1", "risk_text": "Interference unquantified."}
    assert cob.extract_ground_truth_text(item, "risks") == (
        "Interference unquantified."
    )


def test_extract_ground_truth_text_technical_parameter_object() -> None:
    """technical_parameters dict resolves on ``parameter_name``."""
    item = {
        "param_id": "param-1",
        "parameter_name": "7 GHz downlink threshold",
        "value": "minus 47 dBm per megahertz",
        "unit": "dBm/MHz",
    }
    assert cob.extract_ground_truth_text(
        item, "technical_parameters"
    ) == "7 GHz downlink threshold"


def test_extract_ground_truth_text_unknown_dict_first_string() -> None:
    """A dict with NO known field -> the first string value, no halt."""
    item = {"odd_field": "the only text here", "n": 3, "flag": True}
    assert cob.extract_ground_truth_text(item, "decisions") == (
        "the only text here"
    )


def test_extract_ground_truth_text_no_string_falls_back_to_str() -> None:
    """A dict with zero string content -> str(item); never empty, never
    a halt (the line is still written, nothing silently dropped)."""
    item = {"n": 3, "flag": True, "rationale": None}
    out = cob.extract_ground_truth_text(item, "decisions")
    assert out == str(item)
    assert out  # never empty


def test_item_data_always_present_and_lossless(tmp_path: Path) -> None:
    """item_data is always a JSON object holding the FULL original item:
    the dict verbatim for an object, ``{"text": item}`` for a string."""
    dl = _seed(tmp_path)
    decision_obj = {
        "text": "The TIG approved the 7 GHz downlink threshold.",
        "verb": "decided",
        "stakeholders": ["DoD"],
        "confidence": 0.85,
        "rationale": None,
    }
    response = json.dumps(
        {
            "decisions": [
                "A plain string decision.",
                decision_obj,
            ],
            "action_items": [],
            "open_questions": [],
            "commitments": [],
            "risks": [],
            "cross_references": [],
            "attendees": [],
            "topics": [],
            "regulatory_references": [],
            "technical_parameters": [],
            "named_artifacts": [],
            "scheduled_events": [],
        }
    )
    cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(response),
    )
    rows = _written_records(dl)
    assert len(rows) == 2
    by_text = {r["ground_truth_text"]: r for r in rows}
    # String item: wrapped, original recoverable.
    assert by_text["A plain string decision."]["item_data"] == {
        "text": "A plain string decision."
    }
    # Structured item: stored verbatim, every field preserved.
    obj_row = by_text[
        "The TIG approved the 7 GHz downlink threshold."
    ]
    assert obj_row["item_data"] == decision_obj
    for r in rows:
        assert "item_data" in r and isinstance(r["item_data"], dict)


def test_structured_decisions_object_does_not_halt(
    tmp_path: Path,
) -> None:
    """Regression for the reported bug: an all-structured-object
    response (the shape the canonical prompt asks Opus for) writes the
    baseline instead of halting malformed_llm_response."""
    dl = _seed(tmp_path)
    response = json.dumps(
        {
            "decisions": [
                {
                    "text": "Approved the 7 GHz downlink threshold.",
                    "verb": "approved",
                    "stakeholders": ["DoD", "NTIA"],
                    "confidence": 0.9,
                    "rationale": "PCC directed it.",
                }
            ],
            "action_items": [
                {"text": "NTIA to circulate methodology.", "owner": "NTIA"}
            ],
            "open_questions": [
                {"question_text": "Coordination distance?"}
            ],
        }
    )
    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(response),
    )
    assert result["status"] == "success"
    rows = _written_records(dl)
    texts = sorted(r["ground_truth_text"] for r in rows)
    assert texts == [
        "Approved the 7 GHz downlink threshold.",
        "Coordination distance?",
        "NTIA to circulate methodology.",
    ]


# --------------------------------------------------------------------------
# 9. UUID5 pair_ids deterministic across two SEPARATE calls
# --------------------------------------------------------------------------
def test_uuid5_pair_ids_deterministic(tmp_path: Path) -> None:
    def run(root: Path) -> list[str]:
        dl = _seed(root)
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=SpyClient(_VALID_RESPONSE),
        )
        return [
            json.loads(ln)["pair_id"]
            for ln in _out_path(dl).read_text(
                encoding="utf-8"
            ).splitlines()
            if ln.strip()
        ]

    first = run(tmp_path / "a")
    second = run(tmp_path / "b")
    assert first == second, "identical input must yield identical pair_ids"
    # Sanity: ids are real, distinct UUID5 values.
    assert len(set(first)) == len(first)
    for pid in first:
        assert uuid.UUID(pid).version == 5


# --------------------------------------------------------------------------
# 10. The prompt passed to the LLM IS the canonical file (no substitution)
# --------------------------------------------------------------------------
def test_prompt_passed_matches_canonical_file(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    spy = SpyClient(_VALID_RESPONSE)
    cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=True,
        skip_existing=True,
        model=MODEL,
        client=spy,
    )
    canonical = cob._PROMPT_PATH.read_text(encoding="utf-8")
    assert spy.last_system == canonical, (
        "the system prompt sent to the model must be the verbatim "
        "meeting_minutes_llm.md — no silent substitution"
    )
    # And it is the SAME file the Haiku pipeline workflow loads.
    from spectrum_systems_core.workflows import meeting_minutes_llm
    assert cob._PROMPT_PATH == meeting_minutes_llm._PROMPT_PATH


# --------------------------------------------------------------------------
# Model string discipline: --model required, never hardcoded
# --------------------------------------------------------------------------
def test_empty_model_halts(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=True,
            skip_existing=True,
            model="   ",
            client=SpyClient(_VALID_RESPONSE),
        )
    assert ei.value.reason == "missing_model"


def test_no_hardcoded_model_string_in_script() -> None:
    """Zero ``claude-<tier>-<n>`` literals in the script body."""
    import re

    src = (
        SCRIPTS_DIR / "create_opus_reference_baselines.py"
    ).read_text(encoding="utf-8")
    hits = re.findall(r"claude-(?:opus|sonnet|haiku|mythos)-\d[\w-]*", src)
    assert not hits, f"hardcoded model string(s) in script: {hits}"


# --------------------------------------------------------------------------
# Step 4: git-ignored artifact is refused, not silently written
# --------------------------------------------------------------------------
def test_gitignored_artifact_is_refused(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    subprocess.run(
        ["git", "init", "-q", str(dl)], check=True
    )
    # Mirror the real data-lake repo: file-level **/processed/** bulk
    # ignore with NO negation -> the baseline is shadowed.
    (dl / ".gitignore").write_text(
        "**/processed/**\n", encoding="utf-8"
    )
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=SpyClient(_VALID_RESPONSE),
        )
    assert ei.value.reason == "gitignore_blocks_artifact"


def test_gitignore_negation_allows_artifact(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    subprocess.run(["git", "init", "-q", str(dl)], check=True)
    # Mirror exactly what the workflow's "Ensure data-lake .gitignore
    # negation" step writes on top of the data-lake repo's file-glob
    # bulk ignore: re-include the directory chain first (git cannot
    # re-include a file under an excluded parent), then the file.
    (dl / ".gitignore").write_text(
        "**/processed/**\n"
        "!**/processed/**/\n"
        "!**/processed/**/reference_baselines/"
        "opus_reference_minutes.jsonl\n",
        encoding="utf-8",
    )
    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(_VALID_RESPONSE),
    )
    assert result["status"] == "success"
    assert _out_path(dl).is_file()


# --------------------------------------------------------------------------
# source_id filter selects exactly one transcript
# --------------------------------------------------------------------------
def test_source_id_filter(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    _make_docx(
        dl / "store" / "raw" / "transcripts" / "other-transcript-2026.docx"
    )
    other_proc = (
        dl / "store" / "processed" / "meetings"
        / "other-transcript-2026"
    )
    other_proc.mkdir(parents=True)
    (other_proc / "source_record.json").write_text(
        json.dumps(
            make_source_record("other-transcript-2026", str(uuid.uuid4()))
        )
    )
    result = cob.create_baselines(
        data_lake=dl,
        source_id=SOURCE_ID,
        dry_run=True,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(_VALID_RESPONSE),
    )
    assert result["transcripts"] == 1
    assert result["per_transcript"][0]["source_id"] == SOURCE_ID


# --------------------------------------------------------------------------
# Fix 1 — markdown fence stripping runs BEFORE json.loads, always
# --------------------------------------------------------------------------
def _written_records(dl: Path) -> list[dict]:
    return [
        json.loads(ln)
        for ln in _out_path(dl).read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]


def test_parse_response_strips_json_language_fence() -> None:
    """```json\\n{...}\\n``` -> valid JSON parsed (would fail raw)."""
    fenced = "```json\n" + _VALID_RESPONSE + "\n```"
    # Sanity: the fenced blob is NOT itself valid JSON, so a pass here
    # proves the fence was stripped before json.loads (not bypassed).
    with pytest.raises(json.JSONDecodeError):
        json.loads(fenced)
    doc = cob.parse_response(fenced)
    assert isinstance(doc, dict)
    assert doc["decisions"][0].startswith("The group approved")


def test_parse_response_strips_bare_fence_no_language() -> None:
    """``` ... ``` with no language token -> valid JSON parsed."""
    fenced = "```\n" + _VALID_RESPONSE + "\n```"
    with pytest.raises(json.JSONDecodeError):
        json.loads(fenced)
    doc = cob.parse_response(fenced)
    assert isinstance(doc, dict)
    assert "risks" in doc


def test_parse_response_clean_json_still_parses() -> None:
    """Regression: Fix 1 must not break an unfenced clean response."""
    doc = cob.parse_response(_VALID_RESPONSE)
    assert isinstance(doc, dict)
    assert len(doc["decisions"]) == 2


def test_fenced_response_end_to_end_writes_baseline(
    tmp_path: Path,
) -> None:
    """The production path (parse_response_with_recovery) also strips
    fences before json.loads — a ```json fenced response writes the
    full baseline."""
    dl = _seed(tmp_path)
    fenced = "```json\n" + _VALID_RESPONSE + "\n```"
    cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient(fenced),
    )
    assert len(_written_records(dl)) == 8


# --------------------------------------------------------------------------
# Fix 3 Step B — truncation fallback: valid JSON then trailing prose
# --------------------------------------------------------------------------
def test_truncation_fallback_recovers_trailing_prose(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    dl = _seed(tmp_path)
    # Valid object, then the model kept talking (the documented
    # malformed_llm_response root cause: prose after the JSON).
    raw = _VALID_RESPONSE + "\n\nHere is a short summary of the above."
    with pytest.raises(json.JSONDecodeError):
        json.loads(raw)  # sanity: the whole blob is NOT valid JSON
    spy = SpyClient(raw)
    cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=spy,
    )
    # Recovered without a second model call.
    assert spy.calls == 1
    assert len(_written_records(dl)) == 8
    err = capsys.readouterr().err
    assert "truncated_response_used" in err
    assert "char response" in err


def test_truncation_never_accepts_invalid_json(
    tmp_path: Path,
) -> None:
    """A ``}`` inside an unterminated string must NOT yield a
    falsely-accepted object — it must fall through to the retry/halt."""
    dl = _seed(tmp_path)
    # Unterminated string that merely CONTAINS a brace. Truncating at
    # that brace produces invalid JSON; nothing partial may be accepted.
    bad = '{"decisions": ["the rule said {x} and then'
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=SpyClient(bad),  # retry returns same -> halt
        )
    assert ei.value.reason == "malformed_llm_response"
    assert not _out_path(dl).exists()


# --------------------------------------------------------------------------
# Fix 3 Step C — second (simplified) API call recovers clean JSON,
# sending ONLY the first 200 chars of the failed response
# --------------------------------------------------------------------------
def test_second_attempt_recovers_and_sends_only_200_chars(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    dl = _seed(tmp_path)
    # 500-char non-JSON blob with NO brace -> truncation impossible,
    # forcing the Step C retry.
    failed = "x" * 500
    seq = SequenceClient([failed, _VALID_RESPONSE])
    cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=seq,
    )
    assert seq.calls == 2, "must make exactly one repair call"
    assert len(_written_records(dl)) == 8
    err = capsys.readouterr().err
    assert "second_attempt_used" in err

    retry_user = seq.users[1]
    # Red-team: ONLY the first 200 chars of the failed response, never
    # the whole broken blob.
    assert ("x" * 200) in retry_user
    assert ("x" * 201) not in retry_user
    assert failed not in retry_user
    assert "Return ONLY a valid JSON object" in retry_user


def test_both_attempts_fail_halts_not_empty(
    tmp_path: Path,
) -> None:
    """Truncation impossible + retry still malformed -> HALT
    malformed_llm_response. Never a silent empty extraction."""
    dl = _seed(tmp_path)
    seq = SequenceClient(
        ["not json one no brace", "still not json two no brace"]
    )
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=seq,
        )
    assert ei.value.reason == "malformed_llm_response"
    assert seq.calls == 2, "exactly one repair attempt, then halt"
    assert not _out_path(dl).exists(), "no partial/empty JSONL on halt"


def test_non_object_json_halts_without_burning_retry(
    tmp_path: Path,
) -> None:
    """Valid JSON of the wrong shape ([...]) halts immediately — a
    retry cannot turn a list into the required object."""
    dl = _seed(tmp_path)
    seq = SequenceClient(["[1, 2, 3]", _VALID_RESPONSE])
    with pytest.raises(cob.ReferenceBaselineError) as ei:
        cob.create_baselines(
            data_lake=dl,
            source_id=None,
            dry_run=False,
            skip_existing=True,
            model=MODEL,
            client=seq,
        )
    assert ei.value.reason == "malformed_llm_response"
    assert seq.calls == 1, "must NOT waste a repair call on a shape error"


# --------------------------------------------------------------------------
# Red-team: empty {} is valid JSON -> zero items must WARN, not be silent
# --------------------------------------------------------------------------
def test_empty_object_warns_not_silent(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    dl = _seed(tmp_path)
    result = cob.create_baselines(
        data_lake=dl,
        source_id=None,
        dry_run=False,
        skip_existing=True,
        model=MODEL,
        client=SpyClient("{}"),
    )
    assert result["status"] == "success"
    t = result["per_transcript"][0]
    assert t["total"] == 0
    err = capsys.readouterr().err
    assert "empty_extraction" in err, (
        "an empty object on a content-bearing transcript must be loud"
    )


# --------------------------------------------------------------------------
# Fix 4 — registry holds the new Opus string; the workflow resolves it
# from the registry (no hardcoded model literal) so OPUS_MODEL == registry
# --------------------------------------------------------------------------
EXPECTED_OPUS_MODEL = "claude-opus-4-7"


def test_registry_opus_reference_baseline_updated() -> None:
    registry = json.loads(
        (REPO_ROOT / "ai" / "registry" / "model_registry.json").read_text(
            encoding="utf-8"
        )
    )
    assert (
        registry["models"]["opus_reference_baseline"]
        == EXPECTED_OPUS_MODEL
    )


def test_workflow_opus_model_resolves_from_registry() -> None:
    """The workflow must read OPUS_MODEL from the registry key (not a
    literal), so the effective OPUS_MODEL equals the registry entry."""
    import re

    wf = (
        REPO_ROOT
        / ".github"
        / "workflows"
        / "create-opus-reference-baselines.yml"
    ).read_text(encoding="utf-8")
    assert "['models']['opus_reference_baseline']" in wf, (
        "workflow no longer resolves OPUS_MODEL from the registry key"
    )
    model_re = re.compile(r"claude-(?:opus|sonnet|haiku|mythos)-\d[\w-]*")
    for i, line in enumerate(wf.splitlines(), 1):
        stripped = line.strip()
        if stripped.startswith("#"):
            continue
        assert not model_re.search(line), (
            f"hardcoded model literal on non-comment line {i}: {line!r}"
        )


# --------------------------------------------------------------------------
# Fix 2 — max_tokens raised to 16384
# --------------------------------------------------------------------------
def test_opus_max_tokens_constant() -> None:
    """_OPUS_MAX_TOKENS must be 16384 (35% tokenizer headroom for opus-4-7)."""
    assert cob._OPUS_MAX_TOKENS == 16384


# --------------------------------------------------------------------------
# Fix 3 — AnthropicJSONClient strips temperature/top_p/top_k for opus-4-7
# --------------------------------------------------------------------------
def test_opus_47_no_sampling_params_in_api_call(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """For claude-opus-4-7, temperature/top_p/top_k must not reach the SDK."""
    import sys
    import types

    from spectrum_systems_core.workflows import llm_client

    captured: dict = {}

    class _FakeContent:
        text = '{"decisions": []}'

    class _FakeMessage:
        content = [_FakeContent()]

    class _FakeMessages:
        def create(self, **kwargs: object) -> _FakeMessage:
            captured.update(kwargs)
            return _FakeMessage()

    class _FakeAnthropic:
        messages = _FakeMessages()

        def __init__(self) -> None:
            pass

    fake_anthropic = types.ModuleType("anthropic")
    fake_anthropic.Anthropic = _FakeAnthropic  # type: ignore[attr-defined]
    monkeypatch.setitem(sys.modules, "anthropic", fake_anthropic)

    client_obj = llm_client.AnthropicJSONClient(
        model="claude-opus-4-7", max_tokens=16384
    )
    client_obj(system="sys", user="user_text")

    assert "temperature" not in captured, "temperature must not reach the API for opus-4-7"
    assert "top_p" not in captured, "top_p must not reach the API for opus-4-7"
    assert "top_k" not in captured, "top_k must not reach the API for opus-4-7"


# --------------------------------------------------------------------------
# Fix 4 — registry entry has deprecated_params metadata
# --------------------------------------------------------------------------
def test_registry_opus_has_deprecated_params() -> None:
    """Registry must document the deprecated sampling params for opus-4-7."""
    registry = json.loads(
        (REPO_ROOT / "ai" / "registry" / "model_registry.json").read_text(
            encoding="utf-8"
        )
    )
    metadata = registry.get("model_metadata", {}).get("opus_reference_baseline", {})
    assert "deprecated_params" in metadata, (
        "model_metadata.opus_reference_baseline must have deprecated_params"
    )
    assert "temperature" in metadata["deprecated_params"]
    assert "top_p" in metadata["deprecated_params"]
    assert "top_k" in metadata["deprecated_params"]
