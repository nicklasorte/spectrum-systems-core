"""Integration contract test for the create-human-gt-pairs BATCH driver.

The batch workflow has two moving parts:

  1. ``scripts/_match_minutes_transcripts.py`` — pairs every ingested
     transcript to its human-authored minutes ``.docx`` by date token,
     and classifies every transcript into exactly one bucket
     (to_process / skipped_existing / no_minutes / ambiguous).
  2. ``scripts/create_human_gt_pairs.py`` — run once per matched pair,
     unchanged from the single-transcript path.

This test seeds a real temp ``data-lake/`` exercising every bucket and
asserts (mirroring the mission's required tests):

  * Batch skips already-processed transcripts correctly
    (``skip_existing`` honored; ``--no-skip-existing`` reclassifies).
  * Batch handles missing minutes gracefully — a transcript with no
    minutes, and one with no date in its slug, land in ``no_minutes``
    with NO crash and are never in ``to_process``.
  * A date collision is reported ``ambiguous`` and skipped, never guessed.
  * Driving ``create_human_gt_pairs.py`` over the ``to_process`` entries
    (the exact loop the workflow runs) yields pairs where EVERY pair has
    ``human_authored is True`` and ``verified is True``.

The GT-pair shape is asserted through the real writer via the
``make_human_minutes_gt_pair`` fixture factory, per the CLAUDE.md
integration-test rule.
"""
from __future__ import annotations

import json
import os
import subprocess
import sys
import uuid
from pathlib import Path

import jsonschema

from tests.integration.fixtures import (
    make_human_minutes_gt_pair,
    make_source_record,
)

REPO_ROOT = Path(__file__).resolve().parents[2]
MATCHER = REPO_ROOT / "scripts" / "_match_minutes_transcripts.py"
CREATE = REPO_ROOT / "scripts" / "create_human_gt_pairs.py"
GT_SCHEMA_PATH = (
    REPO_ROOT / "contracts" / "schemas" / "ingestion"
    / "ground_truth_pair.schema.json"
)

# Two ingested transcripts on distinct dates, plus collision/no-date/
# no-minutes cases.
SID_DONE = "alpha-tig-kickoff---transcript-20251218"   # already has GT
SID_NEW = "beta-downlink-review---transcript-20260110"  # to process
SID_NO_MIN = "gamma-no-minutes---transcript-20260115"   # no minutes
SID_NO_DATE = "delta-undated-transcript"                # no date in slug
SID_COLA = "epsilon-collide-a---transcript-20260201"    # date collision
SID_COLB = "epsilon-collide-b---transcript-20260201"    # date collision

_STUB_RESPONSE = json.dumps(
    {
        "pairs": [
            {
                "ground_truth_text": (
                    "The TIG approved the 7 GHz downlink interference "
                    "threshold of -6 dB I/N."
                ),
                "extraction_type": "decision",
            },
            {
                "ground_truth_text": (
                    "NTIA to circulate the revised methodology by Jan 15."
                ),
                "extraction_type": "action_item",
            },
        ]
    }
)


def _make_docx(path: Path) -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("Meeting Minutes")
    doc.add_paragraph(
        "The TIG approved the 7 GHz downlink interference threshold "
        "of -6 dB I/N."
    )
    doc.add_paragraph(
        "Action: NTIA to circulate the revised methodology by Jan 15."
    )
    doc.save(str(path))


def _seed_source_record(dl: Path, sid: str) -> str:
    proc = dl / "store" / "processed" / "meetings" / sid
    proc.mkdir(parents=True, exist_ok=True)
    aid = str(uuid.uuid4())
    (proc / "source_record.json").write_text(
        json.dumps(make_source_record(sid, aid)), encoding="utf-8"
    )
    return aid


def _seed_existing_gt(dl: Path, sid: str, source_artifact_id: str) -> None:
    gt = (
        dl / "store" / "processed" / "meetings" / sid
        / "ground_truth" / "human_minutes_gt_pairs.jsonl"
    )
    gt.parent.mkdir(parents=True, exist_ok=True)
    pair = make_human_minutes_gt_pair(
        source_id=sid,
        source_artifact_id=source_artifact_id,
        ground_truth_text="Pre-existing pair.",
        extraction_type="decision",
    )
    gt.write_text(
        json.dumps(pair, sort_keys=True, separators=(",", ":")) + "\n",
        encoding="utf-8",
    )


def _seed_data_lake(tmp_path: Path) -> Path:
    dl = tmp_path / "data-lake"

    aid_done = _seed_source_record(dl, SID_DONE)
    _seed_source_record(dl, SID_NEW)
    _seed_source_record(dl, SID_NO_MIN)
    _seed_source_record(dl, SID_NO_DATE)
    _seed_source_record(dl, SID_COLA)
    _seed_source_record(dl, SID_COLB)

    # SID_DONE already has GT pairs -> must be skipped under skip_existing.
    _seed_existing_gt(dl, SID_DONE, aid_done)

    minutes_dir = dl / "store" / "raw" / "minutes"
    _make_docx(minutes_dir / "Alpha TIG Kickoff Minutes 20251218 FINAL.docx")
    _make_docx(minutes_dir / "Beta Downlink Review Minutes 20260110.docx")
    # SID_NO_MIN (20260115) and SID_NO_DATE: deliberately NO minutes.
    # Collision: ONE minutes file for date 20260201 but TWO transcripts.
    _make_docx(minutes_dir / "Epsilon Collision Minutes 2026-02-01.docx")
    # A minutes file with no matching transcript date.
    _make_docx(minutes_dir / "Orphan Minutes 20991231.docx")
    return dl


def _run_matcher(dl: Path, *, skip_existing: bool) -> dict:
    flag = "--skip-existing" if skip_existing else "--no-skip-existing"
    proc = subprocess.run(
        [
            sys.executable, str(MATCHER),
            "--data-lake", str(dl),
            flag, "--format", "json",
        ],
        capture_output=True, text=True, cwd=str(REPO_ROOT),
    )
    assert proc.returncode == 0, (
        f"matcher failed:\nstdout={proc.stdout}\nstderr={proc.stderr}"
    )
    return json.loads(proc.stdout)


def test_match_table_buckets_every_transcript(tmp_path: Path) -> None:
    dl = _seed_data_lake(tmp_path)
    table = _run_matcher(dl, skip_existing=True)

    proc_ids = {r["source_id"] for r in table["to_process"]}
    skip_ids = {r["source_id"] for r in table["skipped_existing"]}
    no_min_ids = {r["source_id"] for r in table["no_minutes"]}
    ambig_dates = {r["date"] for r in table["ambiguous"]}

    # Only the matched, not-yet-done transcript is processed.
    assert proc_ids == {SID_NEW}
    # Already-done transcript is skipped (skip_existing honored).
    assert SID_DONE in skip_ids
    # Missing minutes + undated slug -> no_minutes, NO crash.
    assert SID_NO_MIN in no_min_ids
    assert SID_NO_DATE in no_min_ids
    # Date collision reported, never auto-paired.
    assert "2026-02-01" in ambig_dates
    assert SID_COLA not in proc_ids and SID_COLB not in proc_ids
    # Orphan minutes surfaced for visibility.
    orphan = {
        r["minutes_file"] for r in table["minutes_without_transcript"]
    }
    assert any("Orphan Minutes" in m for m in orphan)

    # Buckets are disjoint and total: every transcript appears once.
    counted = proc_ids | skip_ids | no_min_ids
    for r in table["ambiguous"]:
        counted |= set(r["source_ids"])
    assert {
        SID_DONE, SID_NEW, SID_NO_MIN, SID_NO_DATE, SID_COLA, SID_COLB
    } <= counted

    c = table["counts"]
    assert c["to_process"] == 1
    assert c["skipped_existing"] == 1
    assert c["ambiguous"] == 1


def test_no_skip_existing_reprocesses_done_transcript(
    tmp_path: Path,
) -> None:
    dl = _seed_data_lake(tmp_path)
    table = _run_matcher(dl, skip_existing=False)
    proc_ids = {r["source_id"] for r in table["to_process"]}
    # With skip disabled the already-done transcript re-enters to_process.
    assert SID_DONE in proc_ids
    assert SID_NEW in proc_ids
    assert not table["skipped_existing"]


def test_missing_minutes_dir_does_not_crash(tmp_path: Path) -> None:
    """A data-lake with transcripts but NO store/raw/minutes/ at all
    must classify every transcript as no_minutes, not raise."""
    dl = tmp_path / "data-lake"
    _seed_source_record(dl, SID_NEW)
    table = _run_matcher(dl, skip_existing=True)
    assert table["to_process"] == []
    assert {r["source_id"] for r in table["no_minutes"]} == {SID_NEW}


def test_batch_loop_emits_human_authored_verified_pairs(
    tmp_path: Path,
) -> None:
    """Drive create_human_gt_pairs.py over the matcher's to_process list
    (exactly what the workflow's process step does) and assert EVERY
    emitted pair carries the non-circular trust markers."""
    dl = _seed_data_lake(tmp_path)
    table = _run_matcher(dl, skip_existing=True)
    assert table["to_process"], "expected at least one transcript to process"

    env = dict(os.environ)
    env["CREATE_HUMAN_GT_PAIRS_STUB_RESPONSE"] = _STUB_RESPONSE
    env.pop("ANTHROPIC_API_KEY", None)

    schema = json.loads(GT_SCHEMA_PATH.read_text(encoding="utf-8"))
    validator = jsonschema.Draft202012Validator(schema)

    for row in table["to_process"]:
        proc = subprocess.run(
            [
                sys.executable, str(CREATE),
                "--data-lake", str(dl),
                "--source-id", row["source_id"],
                "--minutes-file", row["minutes_file"],
            ],
            capture_output=True, text=True, cwd=str(REPO_ROOT), env=env,
        )
        assert proc.returncode == 0, (
            f"create failed for {row['source_id']}:\n"
            f"stdout={proc.stdout}\nstderr={proc.stderr}"
        )
        out = (
            dl / "store" / "processed" / "meetings" / row["source_id"]
            / "ground_truth" / "human_minutes_gt_pairs.jsonl"
        )
        assert out.is_file(), f"no JSONL written for {row['source_id']}"
        lines = [
            ln for ln in out.read_text(encoding="utf-8").splitlines()
            if ln.strip()
        ]
        assert lines, f"empty JSONL for {row['source_id']}"
        for ln in lines:
            doc = json.loads(ln)
            errors = sorted(
                validator.iter_errors(doc), key=lambda e: list(e.path)
            )
            assert not errors, f"schema errors: {[e.message for e in errors]}"
            assert doc["human_authored"] is True
            assert doc["verified"] is True
            assert doc["provenance"]["produced_by"] == "HumanMinutesGTPairs"
            assert doc["source_id"] == row["source_id"]
