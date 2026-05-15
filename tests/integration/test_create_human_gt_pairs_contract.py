"""Integration contract test for ``scripts/create_human_gt_pairs.py``.

This script produces the FIRST non-circular ground truth for the
single-transcript validation baseline: it extracts decisions / action
items / claims from the human-authored minutes ``.docx``, never from
pipeline output. This test:

  1. Uses ``tests.integration.fixtures`` factories that call the REAL
     writers (``make_source_record`` for the identity record,
     ``make_human_minutes_gt_pair`` -> ``create_human_gt_pairs.build_pair``
     for the GT-pair shape assertion).
  2. Seeds a real temp ``data-lake/`` with a real ``.docx`` minutes
     file (built with python-docx) and a ``source_record.json``.
  3. Calls ``create_human_gt_pairs.py`` via ``subprocess.run`` with the
     offline stub seam (no API key needed).
  4. Asserts:
       a. The script exits 0 and writes the JSONL at the contract path.
       b. Every line validates against the ``ground_truth_pair`` schema.
       c. ``human_authored`` is true on every pair (the trust property
          that distinguishes a non-circular pair from a self-referential
          one).
       d. ``provenance.produced_by == "HumanMinutesGTPairs"`` and no
          pair carries pipeline-output provenance.
       e. The script succeeds with NO ``meeting_extraction`` on disk —
          proving it does not read pipeline output.
       f. ``--dry-run`` writes nothing but reports the extracted pairs.
       g. ``annotate_rubric.list_candidates`` can read the JSONL via
          the new ``--gt-file`` seam.
"""
from __future__ import annotations

import json
import subprocess
import sys
import uuid
from pathlib import Path

import jsonschema

from tests.integration.fixtures import (
    make_human_minutes_gt_pair,
    make_source_record,
)

REPO_ROOT = Path(__file__).resolve().parents[2]
SOURCE_ID = "7-ghz-downlink-tig-meeting-kickoff---transcript-20251218"
GT_SCHEMA_PATH = (
    REPO_ROOT / "contracts" / "schemas" / "ingestion"
    / "ground_truth_pair.schema.json"
)
SCRIPT = REPO_ROOT / "scripts" / "create_human_gt_pairs.py"

# What the model would return from the minutes text. The integration
# test never calls Anthropic — the stub env var feeds this verbatim.
_STUB_RESPONSE = json.dumps(
    {
        "pairs": [
            {
                "ground_truth_text": (
                    "The TIG approved the 7 GHz downlink interference "
                    "threshold of -6 dB I/N."
                ),
                "extraction_type": "decision",
            },
            {
                "ground_truth_text": (
                    "NTIA to circulate the revised aggregate "
                    "interference methodology by January 15."
                ),
                "extraction_type": "action_item",
            },
            {
                "ground_truth_text": (
                    "Adjacent-band FSS systems operate with 12 dB of "
                    "additional isolation."
                ),
                "extraction_type": "claim",
            },
        ]
    }
)


def _make_docx(path: Path) -> None:
    """Write a minimal real .docx so DocxExtractor has something to read."""
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("7 GHz Downlink TIG Kickoff Meeting Minutes")
    doc.add_paragraph(
        "The TIG approved the 7 GHz downlink interference threshold "
        "of -6 dB I/N."
    )
    doc.add_paragraph(
        "Action: NTIA to circulate the revised aggregate interference "
        "methodology by January 15."
    )
    doc.save(str(path))


def _seed_data_lake(tmp_path: Path) -> tuple[Path, str, Path]:
    dl = tmp_path / "data-lake"
    artifact_id = str(uuid.uuid4())

    proc_dir = dl / "store" / "processed" / "meetings" / SOURCE_ID
    proc_dir.mkdir(parents=True)
    (proc_dir / "source_record.json").write_text(
        json.dumps(make_source_record(SOURCE_ID, artifact_id))
    )

    minutes = (
        dl / "store" / "raw" / "minutes"
        / "7 GHz Downlink TIG Kickoff Meeting Minutes 20251218 FINAL.docx"
    )
    _make_docx(minutes)
    return dl, artifact_id, minutes


def _run(args: list[str], extra_env: dict | None = None):
    import os

    env = dict(os.environ)
    env[
        "CREATE_HUMAN_GT_PAIRS_STUB_RESPONSE"
    ] = _STUB_RESPONSE
    env.pop("ANTHROPIC_API_KEY", None)  # prove no API path is taken
    if extra_env:
        env.update(extra_env)
    return subprocess.run(
        [sys.executable, str(SCRIPT), *args],
        capture_output=True,
        text=True,
        cwd=str(REPO_ROOT),
        env=env,
    )


def test_writes_schema_valid_human_authored_pairs(tmp_path: Path) -> None:
    data_lake, artifact_id, minutes = _seed_data_lake(tmp_path)

    result = _run(
        [
            "--data-lake", str(data_lake),
            "--source-id", SOURCE_ID,
            "--minutes-file", str(minutes),
        ]
    )
    assert result.returncode == 0, (
        f"script failed:\nstdout={result.stdout}\nstderr={result.stderr}"
    )

    out = (
        data_lake / "store" / "processed" / "meetings" / SOURCE_ID
        / "ground_truth" / "human_minutes_gt_pairs.jsonl"
    )
    assert out.is_file(), "JSONL output not written at the contract path"

    lines = [
        ln for ln in out.read_text(encoding="utf-8").splitlines() if ln.strip()
    ]
    assert len(lines) == 3, f"expected 3 pairs, got {len(lines)}"

    schema = json.loads(GT_SCHEMA_PATH.read_text(encoding="utf-8"))
    validator = jsonschema.Draft202012Validator(schema)
    seen_types = set()
    for ln in lines:
        doc = json.loads(ln)
        errors = sorted(validator.iter_errors(doc), key=lambda e: list(e.path))
        assert not errors, f"schema errors: {[e.message for e in errors]}"
        # Non-circularity trust properties.
        assert doc["human_authored"] is True
        assert doc["verified"] is True
        assert doc["verified_by"] == "human_minutes_20251218"
        assert doc["provenance"]["produced_by"] == "HumanMinutesGTPairs"
        assert doc["source_id"] == SOURCE_ID
        assert doc["source_artifact_id"] == artifact_id
        assert doc["ground_truth_text"].strip()
        assert doc["extraction_type"] in ("decision", "action_item", "claim")
        assert doc["target_type"] == doc["extraction_type"]
        seen_types.add(doc["extraction_type"])
    assert seen_types == {"decision", "action_item", "claim"}


def test_succeeds_without_any_pipeline_extraction(tmp_path: Path) -> None:
    """The script must work with NO meeting_extraction on disk.

    This is the core non-circularity proof: if the script depended on
    pipeline output it would fail here. The only seeded inputs are the
    .docx minutes and source_record.json.
    """
    data_lake, _, minutes = _seed_data_lake(tmp_path)
    ext_dir = data_lake / "store" / "artifacts" / "extractions"
    assert not ext_dir.exists(), "test must not seed any extraction artifact"

    result = _run(
        [
            "--data-lake", str(data_lake),
            "--source-id", SOURCE_ID,
            "--minutes-file", str(minutes),
        ]
    )
    assert result.returncode == 0, result.stderr
    assert not ext_dir.exists(), "script must not create extraction artifacts"


def test_dry_run_writes_nothing_but_reports(tmp_path: Path) -> None:
    data_lake, _, minutes = _seed_data_lake(tmp_path)
    result = _run(
        [
            "--data-lake", str(data_lake),
            "--source-id", SOURCE_ID,
            "--minutes-file", str(minutes),
            "--dry-run",
        ]
    )
    assert result.returncode == 0, result.stderr
    out = (
        data_lake / "store" / "processed" / "meetings" / SOURCE_ID
        / "ground_truth" / "human_minutes_gt_pairs.jsonl"
    )
    assert not out.exists(), "dry-run must not write the JSONL"
    summary = json.loads(result.stdout)
    assert summary["dry_run"] is True
    assert summary["pairs_extracted"] == 3
    assert summary["by_type"] == {
        "decision": 1, "action_item": 1, "claim": 1
    }


def test_fails_loudly_without_source_record(tmp_path: Path) -> None:
    dl = tmp_path / "data-lake"
    minutes = dl / "store" / "raw" / "minutes" / "m.docx"
    _make_docx(minutes)
    (dl / "store" / "processed" / "meetings" / SOURCE_ID).mkdir(parents=True)

    result = _run(
        [
            "--data-lake", str(dl),
            "--source-id", SOURCE_ID,
            "--minutes-file", str(minutes),
        ]
    )
    assert result.returncode != 0
    assert "missing_source_record" in result.stdout


def test_idempotent_pair_ids(tmp_path: Path) -> None:
    data_lake, _, minutes = _seed_data_lake(tmp_path)
    args = [
        "--data-lake", str(data_lake),
        "--source-id", SOURCE_ID,
        "--minutes-file", str(minutes),
    ]
    _run(args)
    out = (
        data_lake / "store" / "processed" / "meetings" / SOURCE_ID
        / "ground_truth" / "human_minutes_gt_pairs.jsonl"
    )
    first = out.read_text(encoding="utf-8")
    _run(args)
    second = out.read_text(encoding="utf-8")
    assert first == second, "re-run must be byte-identical (idempotent)"


def test_fixture_factory_matches_schema() -> None:
    """The fixture factory output must itself be schema-valid so other
    integration tests can rely on it."""
    pair = make_human_minutes_gt_pair(
        source_id=SOURCE_ID,
        source_artifact_id="src-artifact-123",
        ground_truth_text="The TIG deferred the methodology question.",
        extraction_type="decision",
    )
    schema = json.loads(GT_SCHEMA_PATH.read_text(encoding="utf-8"))
    jsonschema.Draft202012Validator(schema).validate(pair)
    assert pair["human_authored"] is True
    assert pair["provenance"]["produced_by"] == "HumanMinutesGTPairs"


def test_annotate_rubric_reads_jsonl_via_gt_file(tmp_path: Path) -> None:
    """The new ``--gt-file`` seam must let annotate_rubric consume the
    human_minutes JSONL (the exact path the updated mobile workflows
    take)."""
    data_lake, _, minutes = _seed_data_lake(tmp_path)
    result = _run(
        [
            "--data-lake", str(data_lake),
            "--source-id", SOURCE_ID,
            "--minutes-file", str(minutes),
        ]
    )
    assert result.returncode == 0, result.stderr
    gt_file = (
        data_lake / "store" / "processed" / "meetings" / SOURCE_ID
        / "ground_truth" / "human_minutes_gt_pairs.jsonl"
    )

    scripts_dir = REPO_ROOT / "scripts"
    if str(scripts_dir) not in sys.path:
        sys.path.insert(0, str(scripts_dir))
    import annotate_rubric  # type: ignore  # noqa: WPS433

    sdl_root = data_lake / "store" / "artifacts"
    pairs = annotate_rubric.list_candidates(
        sdl_root, source_id=SOURCE_ID, limit=50, gt_file=gt_file,
    )
    # Only the decision pair survives the target_type=="decision" filter.
    assert pairs, "annotate_rubric must read the JSONL via --gt-file"
    assert all(p["human_authored"] is True for p in pairs)
