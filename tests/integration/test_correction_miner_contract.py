"""Integration contract test for ``scripts/correction_miner.py``.

CLAUDE.md non-negotiable: a script that reads a pipeline artifact (the
``comparison_result`` envelope) and calls ``validate_artifact`` must
have an integration test that

  1. Uses ``tests/integration/fixtures.py`` factories — never a
     hand-rolled dict — to produce the upstream artifacts via the real
     writers (the real LLM loop, the real Opus-baseline builder, and
     the real ``compare_opus_haiku`` script which writes the
     ``comparison_result`` the miner reads).
  2. Writes artifacts to a real temp directory.
  3. Calls the miner via ``subprocess.run`` against that temp dir.
  4. Asserts the post-conditions (analysis driven by the on-disk
     comparison_result; no PR side effects on a non-promotion).

Opus generation and Haiku evaluation use the explicit offline env-var
transport seams so CI needs no API key. The non-promotion path
guarantees the real ``_default_pr_opener`` (git/gh) is never invoked.
"""
from __future__ import annotations

import json
import os
import subprocess
import sys
import uuid
from pathlib import Path

from tests.integration.fixtures import (
    make_opus_reference_baseline,
    make_promoted_meeting_minutes_artifact,
    make_source_record,
)

REPO_ROOT = Path(__file__).resolve().parents[2]
COMPARE = REPO_ROOT / "scripts" / "compare_opus_haiku.py"
MINER = REPO_ROOT / "scripts" / "correction_miner.py"
SOURCE_ID = "7-ghz-downlink-tig-meeting-kickoff---transcript-20251218"
OPUS_MODEL = "claude-opus-4-6"

# Haiku reproduces only the action item — both decisions become false
# negatives, giving the miner real patterns to mine.
DECISIONS = [
    "The group deferred the aggregate interference methodology.",
    "The group landed on the minus 47 dBm per megahertz threshold.",
]
ACTION_ITEMS = [{"action": "DoD will submit revised ERP values before next session."}]
OPEN_QUESTIONS: list[str] = []


def _seed(tmp_path: Path) -> Path:
    dl = tmp_path / "data-lake"
    store = dl / "store"
    sid_dir = store / "processed" / "meetings" / SOURCE_ID
    sid_dir.mkdir(parents=True)
    source_artifact_id = str(uuid.uuid4())
    (sid_dir / "source_record.json").write_text(
        json.dumps(make_source_record(SOURCE_ID, source_artifact_id)),
        encoding="utf-8",
    )
    # Opus baseline carries the two decisions + the action item.
    make_opus_reference_baseline(
        data_lake_root=dl,
        source_id=SOURCE_ID,
        source_artifact_id=source_artifact_id,
        model=OPUS_MODEL,
        items_by_type={
            "decisions": DECISIONS,
            "action_items": ACTION_ITEMS,
            "open_questions": [],
        },
    )
    # Promoted Haiku artifact reproduces ONLY the action item.
    make_promoted_meeting_minutes_artifact(
        lake_root=store,
        source_id=SOURCE_ID,
        decisions=[],
        action_items=ACTION_ITEMS,
        open_questions=["placeholder question kept non-empty"],
    )
    # A raw transcript so evaluate_candidate has a real input.
    tdir = store / "raw" / "transcripts"
    tdir.mkdir(parents=True)
    (tdir / f"{SOURCE_ID}.txt").write_text(
        "7 GHz Downlink TIG\n"
        + "\n".join(DECISIONS + [a['action'] for a in ACTION_ITEMS])
        + "\n",
        encoding="utf-8",
    )
    return dl


def _seed_transcript_in_processed_only(tmp_path: Path) -> Path:
    """Same as ``_seed`` but the transcript is a real ``.docx`` written
    into ``processed/meetings/<sid>/`` (where transcripts actually
    live) and NOTHING is written under ``raw/transcripts/``. Proves the
    resolver finds the processed-dir transcript end-to-end via the
    subprocess CLI — the exact failure the fix targets."""
    dl = tmp_path / "data-lake"
    store = dl / "store"
    sid_dir = store / "processed" / "meetings" / SOURCE_ID
    sid_dir.mkdir(parents=True)
    source_artifact_id = str(uuid.uuid4())
    (sid_dir / "source_record.json").write_text(
        json.dumps(make_source_record(SOURCE_ID, source_artifact_id)),
        encoding="utf-8",
    )
    make_opus_reference_baseline(
        data_lake_root=dl,
        source_id=SOURCE_ID,
        source_artifact_id=source_artifact_id,
        model=OPUS_MODEL,
        items_by_type={
            "decisions": DECISIONS,
            "action_items": ACTION_ITEMS,
            "open_questions": [],
        },
    )
    make_promoted_meeting_minutes_artifact(
        lake_root=store,
        source_id=SOURCE_ID,
        decisions=[],
        action_items=ACTION_ITEMS,
        open_questions=["placeholder question kept non-empty"],
    )
    # Transcript as a real .docx in the processed dir; no raw/ file.
    from docx import Document

    doc = Document()
    doc.add_paragraph("7 GHz Downlink TIG")
    for line in DECISIONS + [a['action'] for a in ACTION_ITEMS]:
        doc.add_paragraph(line)
    doc.save(str(sid_dir / "transcript.docx"))
    return dl


def _seed_transcript_via_source_record(tmp_path: Path) -> Path:
    """The transcript lives ONLY at ``store/raw/meetings/<sid>/source.txt``
    — NOT co-located with the processed artifacts and NOT under
    ``raw/transcripts/``. ``source_record.json`` (built by the real
    ``make_source_record`` factory) carries ``payload.raw_path``
    pointing at it, relative to the data-lake ``store/`` root. This is
    the exact production layout the #187 glob could not resolve; only
    the source_record step can find the transcript here."""
    dl = tmp_path / "data-lake"
    store = dl / "store"
    sid_dir = store / "processed" / "meetings" / SOURCE_ID
    sid_dir.mkdir(parents=True)
    source_artifact_id = str(uuid.uuid4())
    rel_raw_path = f"raw/meetings/{SOURCE_ID}/source.txt"
    (sid_dir / "source_record.json").write_text(
        json.dumps(
            make_source_record(
                SOURCE_ID, source_artifact_id, raw_path=rel_raw_path
            )
        ),
        encoding="utf-8",
    )
    make_opus_reference_baseline(
        data_lake_root=dl,
        source_id=SOURCE_ID,
        source_artifact_id=source_artifact_id,
        model=OPUS_MODEL,
        items_by_type={
            "decisions": DECISIONS,
            "action_items": ACTION_ITEMS,
            "open_questions": [],
        },
    )
    make_promoted_meeting_minutes_artifact(
        lake_root=store,
        source_id=SOURCE_ID,
        decisions=[],
        action_items=ACTION_ITEMS,
        open_questions=["placeholder question kept non-empty"],
    )
    transcript = store / "raw" / "meetings" / SOURCE_ID / "source.txt"
    transcript.parent.mkdir(parents=True)
    transcript.write_text(
        "7 GHz Downlink TIG\n"
        + "\n".join(DECISIONS + [a['action'] for a in ACTION_ITEMS])
        + "\n",
        encoding="utf-8",
    )
    return dl


def _run(script: Path, args: list[str], extra_env: dict | None = None):
    env = dict(os.environ)
    env.pop("ANTHROPIC_API_KEY", None)
    if extra_env:
        env.update(extra_env)
    return subprocess.run(
        [sys.executable, str(script), *args],
        capture_output=True,
        text=True,
        cwd=str(REPO_ROOT),
        env=env,
    )


def _make_comparison(dl: Path) -> None:
    r = _run(COMPARE, ["--data-lake", str(dl), "--source-id", SOURCE_ID])
    assert r.returncode == 0, f"compare failed: {r.stdout}\n{r.stderr}"


def test_dry_run_reads_validated_comparison_and_analyzes(
    tmp_path: Path,
) -> None:
    dl = _seed(tmp_path)
    _make_comparison(dl)

    r = _run(
        MINER,
        ["--data-lake", str(dl), "--source-id", SOURCE_ID, "--dry-run"],
        extra_env={
            "CORRECTION_MINER_OPUS_STUB_RESPONSE": (
                "ADDITION: when a deferral IS the decision, record it."
            )
        },
    )
    assert r.returncode == 0, f"{r.stdout}\n{r.stderr}"
    out = json.loads(r.stdout)
    assert out["dry_run"] is True
    # Patterns were mined from the real on-disk comparison_result.
    assert out["patterns"], "no patterns mined from comparison_result"
    assert out["candidates"], "no candidates generated"
    for c in out["candidates"]:
        # Opus model came from the registry, not hardcoded.
        assert c["generated_by"]
    assert out["scores"] == []  # dry-run: no evaluation
    assert out["promotion"]["promoted"] is False


def test_non_promotion_path_opens_no_pr(tmp_path: Path) -> None:
    """Haiku stub extracts nothing -> F1 0.0 < baseline -> negative
    delta -> NO promotion -> the real git/gh pr opener is never hit."""
    dl = _seed(tmp_path)
    _make_comparison(dl)

    r = _run(
        MINER,
        [
            "--data-lake", str(dl), "--source-id", SOURCE_ID,
            "--no-dry-run", "--max-candidates", "1",
        ],
        extra_env={
            "CORRECTION_MINER_OPUS_STUB_RESPONSE": "ADD BLOCK",
            "CORRECTION_MINER_HAIKU_STUB_RESPONSE": json.dumps(
                {
                    "decisions": [],
                    "action_items": [],
                    "open_questions": [],
                }
            ),
        },
    )
    assert r.returncode == 0, f"{r.stdout}\n{r.stderr}"
    out = json.loads(r.stdout)
    assert out["dry_run"] is False
    assert out["scores"], "candidate was not evaluated"
    assert out["promotion"]["promoted"] is False
    # No correction branch was created (pr opener never ran).
    branches = subprocess.run(
        ["git", "branch", "--list", "claude/correction-*"],
        cwd=str(REPO_ROOT),
        capture_output=True,
        text=True,
    ).stdout
    assert branches.strip() == ""


def test_transcript_resolves_from_processed_meetings_dir(
    tmp_path: Path,
) -> None:
    """End-to-end: with the transcript ONLY in
    processed/meetings/<sid>/ as a .docx (no raw/ file), the miner
    resolves it and evaluates the candidate. Before the fix this
    failed with ``missing_transcript`` because only raw/ was searched.
    """
    dl = _seed_transcript_in_processed_only(tmp_path)
    _make_comparison(dl)

    r = _run(
        MINER,
        [
            "--data-lake", str(dl), "--source-id", SOURCE_ID,
            "--no-dry-run", "--max-candidates", "1",
        ],
        extra_env={
            "CORRECTION_MINER_OPUS_STUB_RESPONSE": "ADD BLOCK",
            "CORRECTION_MINER_HAIKU_STUB_RESPONSE": json.dumps(
                {
                    "decisions": [],
                    "action_items": [],
                    "open_questions": [],
                }
            ),
        },
    )
    assert r.returncode == 0, f"{r.stdout}\n{r.stderr}"
    assert "missing_transcript" not in r.stdout
    out = json.loads(r.stdout)
    # A non-empty scores list proves evaluate_candidate ran, which
    # only happens if _load_transcript found the processed-dir docx.
    assert out["scores"], "candidate not evaluated (transcript missing)"
    assert out["promotion"]["promoted"] is False


def test_transcript_resolves_via_source_record_raw_path(
    tmp_path: Path,
) -> None:
    """End-to-end: the transcript is ONLY at the location recorded in
    ``source_record.json::payload.raw_path`` (not co-located with the
    processed artifacts, not in raw/transcripts/). The miner subprocess
    must resolve it via the authoritative source_record step and
    evaluate the candidate. Before the fix this failed with
    ``missing_transcript`` because only the glob/raw paths were tried.
    """
    dl = _seed_transcript_via_source_record(tmp_path)
    _make_comparison(dl)

    r = _run(
        MINER,
        [
            "--data-lake", str(dl), "--source-id", SOURCE_ID,
            "--no-dry-run", "--max-candidates", "1",
        ],
        extra_env={
            "CORRECTION_MINER_OPUS_STUB_RESPONSE": "ADD BLOCK",
            "CORRECTION_MINER_HAIKU_STUB_RESPONSE": json.dumps(
                {
                    "decisions": [],
                    "action_items": [],
                    "open_questions": [],
                }
            ),
        },
    )
    assert r.returncode == 0, f"{r.stdout}\n{r.stderr}"
    assert "missing_transcript" not in r.stdout
    out = json.loads(r.stdout)
    # Non-empty scores prove evaluate_candidate ran, which only happens
    # if _load_transcript resolved the transcript via source_record.
    assert out["scores"], "transcript not resolved via source_record"
    assert out["promotion"]["promoted"] is False


def test_transcript_path_override_via_cli(tmp_path: Path) -> None:
    """End-to-end: --transcript-path is honoured by the subprocess CLI
    and bypasses auto-detection (nothing in raw/, transcript file lives
    at an arbitrary path outside the data-lake layout)."""
    dl = _seed_transcript_in_processed_only(tmp_path)
    # Remove the processed-dir transcript so ONLY the override can
    # supply the input.
    (dl / "store" / "processed" / "meetings" / SOURCE_ID
     / "transcript.docx").unlink()
    _make_comparison(dl)

    explicit = tmp_path / "elsewhere" / "explicit_transcript.txt"
    explicit.parent.mkdir(parents=True)
    explicit.write_text(
        "7 GHz Downlink TIG\n" + "\n".join(DECISIONS + [a['action'] for a in ACTION_ITEMS])
        + "\n",
        encoding="utf-8",
    )

    r = _run(
        MINER,
        [
            "--data-lake", str(dl), "--source-id", SOURCE_ID,
            "--no-dry-run", "--max-candidates", "1",
            "--transcript-path", str(explicit),
        ],
        extra_env={
            "CORRECTION_MINER_OPUS_STUB_RESPONSE": "ADD BLOCK",
            "CORRECTION_MINER_HAIKU_STUB_RESPONSE": json.dumps(
                {
                    "decisions": [],
                    "action_items": [],
                    "open_questions": [],
                }
            ),
        },
    )
    assert r.returncode == 0, f"{r.stdout}\n{r.stderr}"
    assert "missing_transcript" not in r.stdout
    out = json.loads(r.stdout)
    assert out["scores"], "override transcript not used"


def test_no_comparisons_halts(tmp_path: Path) -> None:
    dl = tmp_path / "data-lake"
    (dl / "store" / "processed" / "meetings" / SOURCE_ID).mkdir(
        parents=True
    )
    r = _run(
        MINER,
        ["--data-lake", str(dl), "--source-id", SOURCE_ID, "--dry-run"],
    )
    assert r.returncode == 1
    assert "no_comparisons" in r.stdout
