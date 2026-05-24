"""Integration contract test for ``scripts/create_opus_reference_baselines.py``.

CLAUDE.md non-negotiable: a script that reads a pipeline artifact
(``source_record.json``) must have an integration test that

  1. Uses ``tests.integration.fixtures`` factories (``make_source_record``)
     — never a hand-rolled dict — to produce the artifact in the format
     the pipeline writes.
  2. Writes artifacts to a real temp directory (not mocked).
  3. Calls the script via ``subprocess.run`` against that temp dir.
  4. Asserts the correct output on disk (not just the return code).

The script's source_record read requires ONLY a valid-UUID
``artifact_id``; ``source_id`` comes from ``--source-id`` / the
transcript slug and is never required on the record. The
minimal-contract gate (only artifact_id present) is exercised here at
the subprocess level in addition to the in-process unit tests.

The model transport is the explicit offline env-var seam
(``OPUS_REFERENCE_BASELINE_STUB_RESPONSE``) so CI needs no API key.
"""
from __future__ import annotations

import json
import os
import subprocess
import sys
import uuid
from pathlib import Path

from tests.integration.fixtures import make_source_record

REPO_ROOT = Path(__file__).resolve().parents[2]
SCRIPT = REPO_ROOT / "scripts" / "create_opus_reference_baselines.py"
MODEL = "claude-opus-4-6"
TRANSCRIPT_STEM = "phase-ab-transcript-20251218"
SOURCE_ID = "phase-ab-transcript-20251218"

_STUB_RESPONSE = json.dumps(
    {
        "decisions": [
            "The TIG approved the 7 GHz downlink threshold.",
        ],
        "action_items": [{"action": "NTIA to circulate the revised methodology."}],
        "open_questions": [],
        "commitments": [],
        "risks": [
            {
                "risk_id": "risk-1",
                "risk_text": "Adjacent-band interference is unquantified.",
                "raised_by": "FSS Rep",
                "severity": None,
                "mitigation_mentioned": None,
            }
        ],
        "cross_references": [],
        "attendees": [],
        "topics": [],
        "regulatory_references": [],
        "technical_parameters": [],
        "named_artifacts": [],
        "scheduled_events": [],
    }
)


def _make_docx(path: Path) -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("7 GHz Downlink TIG — 2025-12-18")
    doc.add_paragraph("The TIG approved the 7 GHz downlink threshold.")
    doc.add_paragraph("NTIA to circulate the revised methodology.")
    doc.save(str(path))


def _seed(tmp_path: Path, *, with_source_record: bool = True) -> Path:
    dl = tmp_path / "data-lake"
    _make_docx(
        dl / "store" / "raw" / "transcripts" / f"{TRANSCRIPT_STEM}.docx"
    )
    if with_source_record:
        proc = dl / "store" / "processed" / "meetings" / SOURCE_ID
        proc.mkdir(parents=True)
        # Factory: the format the pipeline writes (real-writer-shaped).
        (proc / "source_record.json").write_text(
            json.dumps(make_source_record(SOURCE_ID, str(uuid.uuid4())))
        )
    return dl


def _run(args: list[str], *, env_extra: dict[str, str] | None = None):
    env = dict(os.environ)
    env["OPUS_REFERENCE_BASELINE_STUB_RESPONSE"] = _STUB_RESPONSE
    env.pop("ANTHROPIC_API_KEY", None)  # prove no real API path is taken
    # Pin the chunk-overlap env so a stray CI value cannot leak into the
    # subprocess and quietly change the stamped chunking_strategy_version.
    env.pop("CHUNK_OVERLAP_TURNS", None)
    if env_extra:
        env.update(env_extra)
    return subprocess.run(
        [sys.executable, str(SCRIPT), *args],
        capture_output=True,
        text=True,
        cwd=str(REPO_ROOT),
        env=env,
    )


def _out(dl: Path) -> Path:
    return (
        dl / "store" / "processed" / "meetings" / SOURCE_ID
        / "reference_baselines" / "opus_reference_minutes.jsonl"
    )


def test_subprocess_writes_reference_only_baselines(
    tmp_path: Path,
) -> None:
    dl = _seed(tmp_path)
    artifact_id = json.loads(
        (
            dl / "store" / "processed" / "meetings" / SOURCE_ID
            / "source_record.json"
        ).read_text(encoding="utf-8")
    )["artifact_id"]

    result = _run(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"]
    )
    assert result.returncode == 0, (
        f"stdout={result.stdout}\nstderr={result.stderr}"
    )

    out = _out(dl)
    assert out.is_file(), "JSONL not written at the contract path"
    lines = [
        json.loads(ln)
        for ln in out.read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]
    assert len(lines) == 3  # 1 decision + 1 action + 1 risk
    for r in lines:
        assert r["model_id"] == MODEL
        assert r["model_authored"] is True
        assert r["human_authored"] is False
        assert r["verified"] is False
        assert r["status"] == "reference_only"
        assert r["provenance"]["produced_by"] == (
            "opus_reference_baseline_workflow"
        )
        assert r["source_artifact_id"] == artifact_id
        assert r["source_id"] == SOURCE_ID
        assert r["meeting_date"] == "2025-12-18"
        assert uuid.UUID(r["pair_id"]).version == 5
        # Phase 2.B: subprocess path with no CHUNK_OVERLAP_TURNS set in
        # the env stamps the no-suffix strategy token, matching
        # compare_opus_haiku.py's default for any baseline row that
        # omits the field.
        assert r["chunking_strategy_version"] == "speaker_turn_v1"
    assert {r["extraction_type"] for r in lines} == {
        "decisions", "action_items", "risks"
    }


def test_factory_calls_real_writer_with_strategy_version(
    tmp_path: Path, monkeypatch
) -> None:
    """No-weakening guard: the
    ``tests/integration/fixtures.make_opus_reference_baseline`` factory
    MUST pass ``chunking_strategy_version_value`` through to the real
    writer (per CLAUDE.md's integration-test fixture-factory rule —
    the factory cannot diverge from the writer's signature).

    If a future change makes the writer's strategy-version arg optional
    again, or stops resolving it from ``CHUNK_OVERLAP_TURNS`` via the
    SSOT helper, this test fires before the downstream comparison /
    correction-miner contract suites break in CI.
    """
    monkeypatch.setenv("CHUNK_OVERLAP_TURNS", "2")
    from tests.integration.fixtures import make_opus_reference_baseline

    dl = _seed(tmp_path)
    aid = str(uuid.uuid4())
    out = make_opus_reference_baseline(
        data_lake_root=dl,
        source_id=SOURCE_ID,
        source_artifact_id=aid,
        model=MODEL,
        items_by_type={"decisions": ["one"]},
    )
    rows = [json.loads(ln) for ln in out.read_text("utf-8").splitlines() if ln.strip()]
    assert rows, "factory must produce at least one row"
    for r in rows:
        assert r["chunking_strategy_version"] == "speaker_turn_v1_overlap2"
        assert r["source_artifact_id"] == aid


def test_subprocess_chunk_overlap_turns_env_stamps_version(
    tmp_path: Path,
) -> None:
    """Phase 2.B: subprocess path with `CHUNK_OVERLAP_TURNS=2` set in
    the env stamps `speaker_turn_v1_overlap2` on every row AND the
    summary dict, so the comparison engine's strategy-version gate
    finds a matched value against an overlap=2 haiku artifact."""
    dl = _seed(tmp_path)
    result = _run(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"],
        env_extra={"CHUNK_OVERLAP_TURNS": "2"},
    )
    assert result.returncode == 0, (
        f"stdout={result.stdout}\nstderr={result.stderr}"
    )
    summary = json.loads(result.stdout)
    assert summary["chunk_overlap_turns"] == "2"
    assert summary["chunking_strategy_version"] == "speaker_turn_v1_overlap2"
    for ln in _out(dl).read_text(encoding="utf-8").splitlines():
        if ln.strip():
            assert (
                json.loads(ln)["chunking_strategy_version"]
                == "speaker_turn_v1_overlap2"
            )


def test_subprocess_dry_run_writes_nothing(tmp_path: Path) -> None:
    dl = _seed(tmp_path)
    result = _run(
        ["--data-lake", str(dl), "--model", MODEL, "--dry-run"]
    )
    assert result.returncode == 0, result.stderr
    assert not _out(dl).exists(), "dry-run must not write the JSONL"
    summary = json.loads(result.stdout)
    assert summary["dry_run"] is True
    assert summary["model"] == MODEL
    assert summary["per_transcript"][0]["status"] == "dry_run"


def test_subprocess_missing_source_record_auto_ingests(
    tmp_path: Path,
) -> None:
    """Root-cause fix (subprocess): a transcript present in the
    data-lake but never ingested (no source_record.json — the gate
    that made this workflow exit 2 on every run before any Opus call)
    is self-healed by the deterministic, LLM-free Stage-1 ingestion,
    then the baseline is written. The source_record on disk is shaped
    by the REAL writer (``SourceLoader``), the strongest possible
    real-writer guarantee.
    """
    dl = _seed(tmp_path, with_source_record=False)
    sr_path = (
        dl / "store" / "processed" / "meetings" / SOURCE_ID
        / "source_record.json"
    )
    assert not sr_path.exists(), "precondition: no source_record yet"

    result = _run(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"]
    )
    assert result.returncode == 0, (
        f"stdout={result.stdout}\nstderr={result.stderr}"
    )

    assert sr_path.is_file(), "Stage-1 ingestion must create the record"
    aid = json.loads(sr_path.read_text(encoding="utf-8"))["artifact_id"]
    assert uuid.UUID(aid)

    out = _out(dl)
    assert out.is_file(), "baseline JSONL not written after self-heal"
    lines = [
        json.loads(ln)
        for ln in out.read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]
    assert lines
    for r in lines:
        assert r["source_artifact_id"] == aid


def test_subprocess_only_artifact_id_source_record_passes(
    tmp_path: Path,
) -> None:
    """Gate (subprocess): a source_record with ONLY a valid-UUID
    artifact_id is accepted and the baseline JSONL is written."""
    dl = _seed(tmp_path, with_source_record=False)
    proc = dl / "store" / "processed" / "meetings" / SOURCE_ID
    proc.mkdir(parents=True)
    aid = str(uuid.uuid4())
    (proc / "source_record.json").write_text(
        json.dumps({"artifact_id": aid}), encoding="utf-8"
    )
    result = _run(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"]
    )
    assert result.returncode == 0, (
        f"stdout={result.stdout}\nstderr={result.stderr}"
    )
    out = _out(dl)
    assert out.is_file(), "JSONL not written for minimal source_record"
    for ln in out.read_text(encoding="utf-8").splitlines():
        if ln.strip():
            assert json.loads(ln)["source_artifact_id"] == aid


def test_subprocess_missing_artifact_id_source_record_halts(
    tmp_path: Path,
) -> None:
    """Gate (subprocess): no artifact_id -> invalid_source_record, no
    partial output."""
    dl = _seed(tmp_path, with_source_record=False)
    proc = dl / "store" / "processed" / "meetings" / SOURCE_ID
    proc.mkdir(parents=True)
    (proc / "source_record.json").write_text(
        json.dumps({"source_id": SOURCE_ID}), encoding="utf-8"
    )
    result = _run(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"]
    )
    assert result.returncode != 0
    assert "invalid_source_record" in result.stdout
    assert not _out(dl).exists()


def test_subprocess_pair_ids_deterministic(tmp_path: Path) -> None:
    def ids(root: Path) -> list[str]:
        dl = _seed(root)
        r = _run(
            [
                "--data-lake", str(dl), "--model", MODEL,
                "--no-skip-existing",
            ]
        )
        assert r.returncode == 0, r.stderr
        return [
            json.loads(ln)["pair_id"]
            for ln in _out(dl).read_text(encoding="utf-8").splitlines()
            if ln.strip()
        ]

    assert ids(tmp_path / "a") == ids(tmp_path / "b")


def _run_with_stub(args: list[str], stub: str):
    """Same as ``_run`` but with a caller-supplied stub response.

    The env stub returns the SAME text on every call, so this drives
    the recovery paths that succeed WITHOUT a differing second response
    (fence-strip, truncation, empty-object). The differing-second-call
    path is covered by the in-process unit suite (a sequence client).
    """
    env = dict(os.environ)
    env["OPUS_REFERENCE_BASELINE_STUB_RESPONSE"] = stub
    env.pop("ANTHROPIC_API_KEY", None)
    return subprocess.run(
        [sys.executable, str(SCRIPT), *args],
        capture_output=True,
        text=True,
        cwd=str(REPO_ROOT),
        env=env,
    )


def test_subprocess_fenced_response_recovered(tmp_path: Path) -> None:
    """Fix 1 at the subprocess boundary: a ```json-fenced response is
    fence-stripped before json.loads and the baseline is written."""
    dl = _seed(tmp_path)
    result = _run_with_stub(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"],
        "```json\n" + _STUB_RESPONSE + "\n```",
    )
    assert result.returncode == 0, (
        f"stdout={result.stdout}\nstderr={result.stderr}"
    )
    out = _out(dl)
    assert out.is_file()
    lines = [
        ln for ln in out.read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]
    assert len(lines) == 3  # 1 decision + 1 action + 1 risk


def test_subprocess_trailing_prose_truncated(tmp_path: Path) -> None:
    """Fix 3 Step B at the subprocess boundary: valid JSON followed by
    prose is truncated back to the last ``}`` and recovered."""
    dl = _seed(tmp_path)
    result = _run_with_stub(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"],
        _STUB_RESPONSE + "\n\nThat concludes the extraction. Thanks!",
    )
    assert result.returncode == 0, (
        f"stdout={result.stdout}\nstderr={result.stderr}"
    )
    assert _out(dl).is_file()
    assert "truncated_response_used" in result.stderr
    lines = [
        ln for ln in _out(dl).read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]
    assert len(lines) == 3


def test_subprocess_empty_object_warns_returncode_zero(
    tmp_path: Path,
) -> None:
    """Red-team at the subprocess boundary: ``{}`` is valid JSON, so
    the run succeeds (rc 0) but MUST warn loudly on stderr — never a
    silent empty baseline."""
    dl = _seed(tmp_path)
    result = _run_with_stub(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"],
        "{}",
    )
    assert result.returncode == 0, (
        f"stdout={result.stdout}\nstderr={result.stderr}"
    )
    assert "empty_extraction" in result.stderr
    summary = json.loads(result.stdout)
    assert summary["per_transcript"][0]["total"] == 0


_STRUCTURED_DECISION = {
    "text": "The TIG approved the 7 GHz downlink threshold.",
    "verb": "approved",
    "stakeholders": ["DoD", "NTIA"],
    "confidence": 0.9,
    "rationale": "The PCC directed it.",
}

_STRUCTURED_STUB_RESPONSE = json.dumps(
    {
        # The exact shape the canonical prompt asks Opus for: a
        # ``decisions`` array whose item is a structured object, not a
        # string. This used to halt malformed_llm_response.
        "decisions": [_STRUCTURED_DECISION],
        "action_items": [
            {"text": "NTIA to circulate the methodology.", "owner": "NTIA"}
        ],
        "open_questions": [{"question_text": "Coordination distance?"}],
        "commitments": [],
        "risks": [],
        "cross_references": [],
        "attendees": [],
        "topics": [],
        "regulatory_references": [],
        "technical_parameters": [
            {
                "param_id": "p-1",
                "parameter_name": "downlink threshold",
                "value": "minus 47 dBm/MHz",
            }
        ],
        "named_artifacts": [],
        "scheduled_events": [],
    }
)


def test_subprocess_structured_object_items_recovered(
    tmp_path: Path,
) -> None:
    """End-to-end at the subprocess boundary: a fully structured-object
    response (the canonical prompt's object form) is written, NOT
    halted, and ``item_data`` keeps every original field verbatim."""
    dl = _seed(tmp_path)
    result = _run_with_stub(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"],
        _STRUCTURED_STUB_RESPONSE,
    )
    assert result.returncode == 0, (
        f"stdout={result.stdout}\nstderr={result.stderr}"
    )
    out = _out(dl)
    assert out.is_file(), "structured-object response must still write"
    rows = [
        json.loads(ln)
        for ln in out.read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]
    # 1 decision + 1 action + 1 question + 1 technical_parameter
    assert len(rows) == 4
    by_type = {r["extraction_type"]: r for r in rows}

    # decisions: structured object resolved on ``text``; item_data is
    # the full object, every field preserved (nothing lost).
    dec = by_type["decisions"]
    assert dec["ground_truth_text"] == _STRUCTURED_DECISION["text"]
    assert dec["item_data"] == _STRUCTURED_DECISION

    # action_items dict -> ``text`` field, owner preserved in item_data.
    act = by_type["action_items"]
    assert act["ground_truth_text"] == "NTIA to circulate the methodology."
    assert act["item_data"]["owner"] == "NTIA"

    # open_questions dict -> ``question_text``.
    assert by_type["open_questions"]["ground_truth_text"] == (
        "Coordination distance?"
    )

    # technical_parameters dict -> ``parameter_name`` (priority list).
    tp = by_type["technical_parameters"]
    assert tp["ground_truth_text"] == "downlink threshold"
    assert tp["item_data"]["value"] == "minus 47 dBm/MHz"

    for r in rows:
        assert r["status"] == "reference_only"
        assert isinstance(r["item_data"], dict)


def test_subprocess_unrecoverable_halts_no_partial(
    tmp_path: Path,
) -> None:
    """Garbage with no recoverable object: fence-strip, truncation, and
    the (same-stub) retry all fail -> halt malformed_llm_response, no
    partial JSONL."""
    dl = _seed(tmp_path)
    result = _run_with_stub(
        ["--data-lake", str(dl), "--model", MODEL, "--no-skip-existing"],
        "I am unable to produce JSON for this transcript.",
    )
    assert result.returncode != 0
    assert "malformed_llm_response" in result.stdout
    assert not _out(dl).exists()
